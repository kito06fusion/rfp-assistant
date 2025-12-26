from __future__ import annotations

import logging
import re
import httpx
import os
import tempfile
from io import BytesIO
from pathlib import Path
from typing import Dict, List, Any, Optional
from datetime import datetime

from backend.models import RequirementsResult, ExtractionResult

logger = logging.getLogger(__name__)
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.section import WD_SECTION_START
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. DOCX generation will be disabled.")

try:
    from PIL import Image
    PIL_AVAILABLE = True
except Exception:
    PIL_AVAILABLE = False
    logger.warning("Pillow (PIL) not available. Image validation will be disabled.")


 #function to clean up mermaid diagram label characters
def _sanitize_mermaid_labels(diagram: str) -> str:
    if not diagram:
        return diagram
    
    diagram = (
        diagram.replace('"', '"')
        .replace('"', '"')
        .replace(''', "'")
        .replace(''', "'")
    )
    
    lines = diagram.split('\n')
    fixed_lines = []
    
    for line in lines:
        def fix_label(match):
            label_content = match.group(1)
            # If it already has quotes, leave it alone
            if (label_content.startswith('"') and label_content.endswith('"')) or \
               (label_content.startswith("'") and label_content.endswith("'")):
                return match.group(0)
            # If it contains parentheses, wrap in quotes
            if '(' in label_content or ')' in label_content:
                # Escape any existing quotes in the label
                escaped = label_content.replace('"', '\\"')
                return f'["{escaped}"]'
            return match.group(0)
        
        # Match [label] patterns
        line = re.sub(r'\[([^\]]+)\]', fix_label, line)
        fixed_lines.append(line)
    
    return '\n'.join(fixed_lines).strip()


#function to check if bytes start with a PNG header
def _is_png_bytes(data: bytes) -> bool:
    return bool(data) and data.startswith(b"\x89PNG\r\n\x1a\n")


#function to remove all content from a docx paragraph
def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


#function to setup default styles for DOCX document
def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15
    
    for level in (1, 2, 3, 4):
        try:
            h = doc.styles[f"Heading {level}"]
            h.font.name = "Calibri"
            h.paragraph_format.space_before = Pt(12 if level <= 2 else 8)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            h.paragraph_format.widow_control = True
            h.paragraph_format.keep_together = True
            if level <= 2:
                h.paragraph_format.space_before = Pt(18)
        except KeyError:
            logger.warning(f"Heading {level} style not found, skipping")
        except Exception as e:
            logger.warning(f"Failed to set heading {level} properties: {e}")
    
    try:
        list_bullet = doc.styles["List Bullet"]
        list_bullet.font.name = "Calibri"
        list_bullet.font.size = Pt(11)
        pf = list_bullet.paragraph_format
        pf.left_indent = Inches(0.25)
        pf.first_line_indent = Inches(-0.25)
    except KeyError:
        logger.warning("List Bullet style not found, will use default")


#function to configure page margins, numbering and footer formatting
def setup_page_formatting(doc, start_page_number: int = 1):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    if start_page_number > 1:
        sect_pr = section._sectPr
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), str(start_page_number))
        sect_pr.append(pg_num_type)
    
    footer = section.footer
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


#function to add a modern front page (logo, title, date, company)
def add_modern_front_page(doc, title: str, project_root: Optional[Path] = None):
    if project_root is None:
        project_root = Path(__file__).parent.parent.parent
    
    logo_path = None
    possible_logo_paths = [
        project_root / "frontend" / "src" / "assets" / "logo-transparent.png",
        project_root / "frontend" / "src" / "assets" / "logo.png",
        project_root / "assets" / "logo-transparent.png",
        project_root / "assets" / "logo.png",
        project_root / "backend" / "assets" / "logo-transparent.png",
        project_root / "backend" / "assets" / "logo.png",
    ]
    
    for path in possible_logo_paths:
        if path.exists():
            logo_path = path
            logger.info(f"Found logo at: {logo_path}")
            break
    
    if logo_path is None:
        logger.warning(f"Logo not found. Tried paths: {[str(p) for p in possible_logo_paths]}")
    
    for _ in range(2):
        doc.add_paragraph()
    
    if logo_path and logo_path.exists():
        try:
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = logo_para.add_run()
            if PIL_AVAILABLE:
                try:
                    img = Image.open(logo_path)
                    width, height = img.size
                    aspect_ratio = width / height
                    logo_height = Inches(1.5)
                    logo_width = Inches(1.5 * aspect_ratio)
                    logger.info(f"Logo dimensions: {width}x{height}, display size: {logo_width} x {logo_height}")
                except Exception as e:
                    logger.warning(f"Failed to get logo dimensions: {e}")
                    logo_width = Inches(3.5)
                    logo_height = Inches(1.5)
            else:
                logo_width = Inches(3.5)
                logo_height = Inches(1.5)
            
            run.add_picture(str(logo_path), width=logo_width, height=logo_height)
            logger.info(f"Successfully added logo to front page")
            
            doc.add_paragraph()
            doc.add_paragraph()
        except Exception as e:
            logger.error(f"Failed to add logo to front page: {e}", exc_info=True)
    else:
        logger.warning("Logo not found or not accessible, skipping logo on front page")
    
    title_length = len(title)
    if title_length > 80:
        font_size = Pt(24)
    elif title_length > 60:
        font_size = Pt(28)
    else:
        font_size = Pt(32)
    
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(12)
    title_para.paragraph_format.space_before = Pt(0)
    
    title_para.paragraph_format.widow_control = False
    title_para.paragraph_format.keep_together = False
    title_para.paragraph_format.keep_with_next = False
    
    p_pr = title_para._element.get_or_add_pPr()
    try:
        wrap = OxmlElement('w:wordWrap')
        wrap.set(qn('w:val'), '0')
        p_pr.append(wrap)
        
        overflow = OxmlElement('w:overflowPunct')
        overflow.set(qn('w:val'), '0')
        p_pr.append(overflow)
        
        if p_pr.find(qn('w:ind')) is not None:
            p_pr.remove(p_pr.find(qn('w:ind')))
    except Exception:
        pass
    
    title_run = title_para.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = font_size
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(26, 84, 144)
    
    try:
        r_pr = title_run._element.get_or_add_rPr()
    except Exception:
        pass
    
    logger.info(f"Added title to front page: '{title[:50]}...' (length: {title_length}, font size: {font_size})")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_run.font.name = "Calibri"
    date_run.font.size = Pt(14)
    date_run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(4):
        doc.add_paragraph()
    
    company_para = doc.add_paragraph()
    company_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company_para.add_run("fusionAIx")
    company_run.font.name = "Calibri"
    company_run.font.size = Pt(20)
    company_run.font.bold = True
    company_run.font.color.rgb = RGBColor(26, 84, 144)
    
    website_para = doc.add_paragraph()
    website_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    website_run = website_para.add_run("www.fusionaix.com")
    website_run.font.name = "Calibri"
    website_run.font.size = Pt(12)
    website_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()


#function to insert a manual table of contents into the DOCX
def add_manual_toc(doc, toc_entries: List[Dict[str, Any]]):
    if not toc_entries:
        para = doc.add_paragraph("No table of contents entries available.")
        return
    
    for entry in toc_entries:
        para = doc.add_paragraph()
        para.style = 'Normal'
        
        pf = para.paragraph_format
        indent_level = (entry.get('level', 1) - 1) * 0.5
        if indent_level > 0:
            pf.left_indent = Inches(indent_level)
        
        text = entry.get('text', '')
        
        run = para.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)

#function to style a table header cell (bold text + background)
def set_table_header_cell(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            try:
                r.font.color.rgb = RGBColor(255, 255, 255)
            except Exception as color_err:
                logger.debug("Failed to set font color for table header: %s", color_err)
    try:
        tc_pr = cell._element.get_or_add_tcPr()
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), '1a5490')
        shading.set(qn('w:val'), 'clear')
        tc_pr.append(shading)
    except Exception as e:
        logger.warning("Failed to set table header color: %s", e)


#function to finalize table formatting for all cells and rows
def finalize_table(table):
    try:
        table.autofit = True
    except Exception as autofit_err:
        logger.debug("Failed to set table autofit: %s", autofit_err)
    
    for row in table.rows:
        for cell in row.cells:
            try:
                tc_pr = cell._element.get_or_add_tcPr()
                v_align = OxmlElement('w:vAlign')
                v_align.set(qn('w:val'), 'top')
                tc_pr.append(v_align)
                
                tc_mar = OxmlElement('w:tcMar')
                for margin_name, margin_val in [('top', '80'), ('right', '80'), ('bottom', '80'), ('left', '80')]:
                    margin = OxmlElement(f'w:{margin_name}')
                    margin.set(qn('w:w'), margin_val)
                    margin.set(qn('w:type'), 'dxa')
                    tc_mar.append(margin)
                tc_pr.append(tc_mar)
            except Exception as e:
                logger.debug("Failed to set table cell formatting: %s", e)

#function to add text with simple bold formatting markers to a paragraph
def _add_formatted_text_to_paragraph(para, text: str):
    if not text:
        return
    
    if para.runs:
        clear_paragraph(para)
    
    parts = []
    last_end = 0
    
    for match in re.finditer(r'\*\*(.*?)\*\*', text):
        if match.start() > last_end:
            parts.append(('normal', text[last_end:match.start()]))
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    if last_end < len(text):
        parts.append(('normal', text[last_end:]))
    
    if not parts:
        parts = [('normal', text)]
    
    for fmt_type, content in parts:
        if content:
            run = para.add_run(content)
            if fmt_type == 'bold':
                run.bold = True

#function to add a bulleted paragraph with formatted runs
def _add_bullet_paragraph(doc, content: str):
    para = doc.add_paragraph()
    
    pf = para.paragraph_format
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.25)
    pf.space_after = Pt(6)
    
    bullet_run = para.add_run('• ')
    bullet_run.font.name = "Calibri"
    bullet_run.font.size = Pt(11)
    
    parts = []
    last_end = 0
    
    for match in re.finditer(r'\*\*(.*?)\*\*', content):
        if match.start() > last_end:
            parts.append(('normal', content[last_end:match.start()]))
        parts.append(('bold', match.group(1)))
        last_end = match.end()
    
    if last_end < len(content):
        parts.append(('normal', content[last_end:]))
    
    if not parts:
        parts = [('normal', content)]
    
    for fmt_type, part_content in parts:
        if part_content:
            run = para.add_run(part_content)
            run.font.name = "Calibri"
            run.font.size = Pt(11)
            if fmt_type == 'bold':
                run.bold = True
    
    return para

#function to start a new table with a styled header row
def _start_table(doc, header_cells: List[str]):
    num_cols = len(header_cells)
    current_table = doc.add_table(rows=1, cols=num_cols)
    table_styles = ['Light Grid Accent 1', 'Grid Table 1 Light', 'Table Grid']
    style_applied = False
    for style_name in table_styles:
        try:
            current_table.style = style_name
            style_applied = True
            logger.debug("Applied table style: %s", style_name)
            break
        except Exception as e:
            logger.debug("Failed to apply table style '%s': %s", style_name, e)
            continue
    
    if not style_applied:
        logger.debug("No table style could be applied, using default")
    
    try:
        header_row = current_table.rows[0]
        tr = header_row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)
    except Exception as e:
        logger.warning(f"Failed to set table header row property: {e}")
    
    header_row_cells = current_table.rows[0].cells
    for col_idx, cell_text in enumerate(header_cells):
        cell = header_row_cells[col_idx]
        cell.text = cell_text
        set_table_header_cell(cell)
    
    return current_table


#function to extract heading lines from markdown text
def _extract_headings_from_markdown(text: str) -> List[Dict[str, Any]]:
    headings = []
    lines = text.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith('#### '):
            headings.append({'text': _clean_markdown_text(stripped[5:]), 'level': 4})
        elif stripped.startswith('### '):
            headings.append({'text': _clean_markdown_text(stripped[4:]), 'level': 3})
        elif stripped.startswith('## '):
            headings.append({'text': _clean_markdown_text(stripped[3:]), 'level': 2})
        elif stripped.startswith('# '):
            headings.append({'text': _clean_markdown_text(stripped[2:]), 'level': 1})
    
    return headings


#function to parse markdown content and insert formatted elements into DOCX
def _parse_markdown_to_docx(doc, text: str):
    if not text:
        return
    
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if re.match(r'^---+$', stripped):
            continue
        cleaned_lines.append(line)
    
    lines = cleaned_lines
    i = 0
    in_table = False
    current_table = None
    in_code_block = False
    code_block_lines = []
    code_block_lang = None
    last_was_blank = False
    list_item_buffer = []
    diagram_count = 0
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        mermaid_start_pattern = re.compile(r'^(?:flowchart|graph|sequenceDiagram|classDiagram|gantt|stateDiagram|pie|erDiagram)\b', re.IGNORECASE)
        if mermaid_start_pattern.match(stripped):
            if diagram_count >= 1:
                logger.info('Skipping additional mermaid diagram (only 1 diagram per document allowed)')
                j = i + 1
                while j < len(lines):
                    next_line = lines[j]
                    next_stripped = next_line.strip()
                    if not next_stripped or next_stripped.startswith('```') or next_stripped.startswith('#'):
                        break
                    j += 1
                i = j
                continue
            
            diagram_count += 1
            block_lines = [line.rstrip('\n')]
            j = i + 1
            caption_text = None
            while j < len(lines):
                next_line = lines[j]
                next_stripped = next_line.strip()
                if not next_stripped:
                    break
                if re.match(r'^Caption:\s*', next_stripped, flags=re.IGNORECASE):
                    caption_text = re.sub(r'^Caption:\s*', '', next_stripped, flags=re.IGNORECASE).strip()
                    j += 1
                    break
                if next_stripped.startswith('```') or next_stripped.startswith('#'):
                    break
                block_lines.append(next_line.rstrip('\n'))
                j += 1

            block_text = '\n'.join(block_lines)
            sanitized_block = _sanitize_mermaid_labels(block_text)

            png_bytes = None
            
            try:
                logger.debug('Attempting Kroki PNG rendering for diagram')
                kroki_url = 'https://kroki.io/mermaid/png'
                resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                if resp.status_code == 200:
                    kroki_png = resp.content
                    if kroki_png and _is_png_bytes(kroki_png):
                        png_bytes = kroki_png
                        logger.info('Kroki PNG rendering succeeded (%d bytes)', len(png_bytes))
                    else:
                        logger.warning('Kroki returned invalid PNG (status 200 but no PNG header, %d bytes)', len(kroki_png) if kroki_png else 0)
                else:
                    logger.warning('Kroki returned status %s for mermaid render', resp.status_code)
                    if resp.status_code == 400:
                        error_msg = resp.text[:500] if hasattr(resp, 'text') and resp.text else 'Unknown error'
                        logger.error('Kroki syntax error - diagram code preview: %s', sanitized_block[:200])
                        logger.error('Kroki error details: %s', error_msg)
            except Exception as e:
                logger.exception('Kroki PNG rendering failed: %s', e)

            img_to_insert = png_bytes if png_bytes else None

            if img_to_insert:
                try:
                    logger.info('Inserting locally rendered mermaid image into DOCX (unfenced)')
                    
                    validated_bytes = img_to_insert
                    if PIL_AVAILABLE and isinstance(img_to_insert, bytes) and len(img_to_insert) > 0:
                        try:
                            img_stream_for_validation = BytesIO(img_to_insert)
                            img_stream_for_validation.seek(0)
                            im = Image.open(img_stream_for_validation)
                            im.verify()
                            img_stream_for_save = BytesIO(img_to_insert)
                            img_stream_for_save.seek(0)
                            im = Image.open(img_stream_for_save)
                            reencoded_stream = BytesIO()
                            im.save(reencoded_stream, format='PNG')
                            reencoded_bytes = reencoded_stream.getvalue()
                            if reencoded_bytes and len(reencoded_bytes) > 0:
                                validated_bytes = reencoded_bytes
                                logger.debug('PNG validated and re-encoded successfully (%d bytes)', len(validated_bytes))
                            else:
                                logger.warning('PNG re-encoding produced empty result, will use original bytes')
                        except Exception as pil_err:
                            logger.warning('PNG validation/re-encoding failed: %s, will use original bytes', str(pil_err))
                    
                    inserted = False
                    try:
                        img_stream = BytesIO(validated_bytes)
                        img_stream.seek(0)
                        doc.add_picture(img_stream, width=Inches(5))
                        inserted = True
                        logger.debug('Successfully inserted PNG from BytesIO')
                    except Exception as stream_err:
                        logger.debug('doc.add_picture from BytesIO failed: %s, trying temp file fallback', str(stream_err))
                    
                    if not inserted:
                        tmpname = None
                        try:
                            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tf:
                                tf.write(validated_bytes)
                                tf.flush()
                                tmpname = tf.name
                            doc.add_picture(tmpname, width=Inches(5))
                            inserted = True
                            logger.debug('Successfully inserted PNG from temp file')
                        except Exception as file_err:
                            logger.warning('Failed to insert PNG from temp file: %s', str(file_err))
                        finally:
                            if tmpname:
                                try:
                                    os.remove(tmpname)
                                except Exception:
                                    pass
                    
                    if inserted and caption_text:
                        cap_para = doc.add_paragraph()
                        cap_run = cap_para.add_run(caption_text)
                        cap_run.italic = True
                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif not inserted:
                        logger.warning('Failed to insert PNG image, will fall back to Kroki')
                        img_to_insert = None
                except Exception as e:
                    logger.exception('Failed to insert locally rendered mermaid image into docx (unfenced): %s', e)
                    img_to_insert = None

            if not img_to_insert:
                try:
                    kroki_url = 'https://kroki.io/mermaid/png'
                    logger.debug('Attempting Kroki fallback for mermaid diagram (length=%d chars)', len(sanitized_block))
                    resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                    if resp.status_code == 200:
                        img_bytes = resp.content
                        if img_bytes and len(img_bytes) > 0:
                            if img_bytes.startswith(b"\x89PNG\r\n\x1a\n"):
                                try:
                                    img_stream = BytesIO(img_bytes)
                                    doc.add_picture(img_stream, width=Inches(5))
                                    if caption_text:
                                        cap_para = doc.add_paragraph()
                                        cap_run = cap_para.add_run(caption_text)
                                        cap_run.italic = True
                                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    logger.info('Successfully inserted Kroki-rendered mermaid diagram')
                                except Exception as e:
                                    logger.warning('Failed to insert Kroki mermaid image into docx (unfenced): %s', e)
                                    para = doc.add_paragraph(style='Normal')
                                    para.style.font.name = 'Consolas'
                                    para.style.font.size = Pt(10)
                                    para.text = block_text
                            else:
                                logger.warning('Kroki returned invalid PNG (no PNG header, %d bytes)', len(img_bytes))
                                para = doc.add_paragraph(style='Normal')
                                para.style.font.name = 'Consolas'
                                para.style.font.size = Pt(10)
                                para.text = block_text
                        else:
                            logger.warning('Kroki returned empty response')
                            para = doc.add_paragraph(style='Normal')
                            para.style.font.name = 'Consolas'
                            para.style.font.size = Pt(10)
                            para.text = block_text
                    else:
                        error_text = resp.text[:500] if hasattr(resp, 'text') else 'no error text'
                        logger.warning('Kroki returned status %s for mermaid render (unfenced): %s', resp.status_code, error_text)
                        logger.debug('Diagram that failed: %s', sanitized_block[:200])
                        para = doc.add_paragraph(style='Normal')
                        para.style.font.name = 'Consolas'
                        para.style.font.size = Pt(10)
                        para.text = block_text
                except Exception as e:
                    logger.exception('Failed to call Kroki for mermaid rendering (unfenced): %s', e)
                    para = doc.add_paragraph(style='Normal')
                    para.style.font.name = 'Consolas'
                    para.style.font.size = Pt(10)
                    para.text = block_text

            i = j
            continue
        
        if stripped.startswith('```'):
            lang = stripped[3:].strip().lower()
            if in_code_block:
                if code_block_lines:
                    block_text = '\n'.join(code_block_lines)
                    caption_text = None
                    m_caption = re.search(r"\n?\s*Caption:\s*(.+)\s*$", block_text, flags=re.IGNORECASE)
                    if m_caption:
                        caption_text = m_caption.group(1).strip()
                        block_text = re.sub(r"\n?\s*Caption:\s*.+\s*$", "", block_text, flags=re.IGNORECASE)

                    if code_block_lang == 'mermaid':
                        if diagram_count >= 1:
                            logger.info('Skipping additional mermaid diagram in fenced block (only 1 diagram per document allowed)')
                            para = doc.add_paragraph(style='Normal')
                            para.style.font.name = 'Consolas'
                            para.style.font.size = Pt(10)
                            para.text = block_text
                        else:
                            diagram_count += 1
                            rendered = None
                            sanitized_block = _sanitize_mermaid_labels(block_text)
                            try:
                                logger.debug('Attempting Kroki PNG rendering for fenced diagram')
                                kroki_url = 'https://kroki.io/mermaid/png'
                                resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                                if resp.status_code == 200:
                                    rendered = resp.content
                                    if rendered and _is_png_bytes(rendered):
                                        logger.info('Kroki PNG rendering succeeded for fenced diagram (%d bytes)', len(rendered))
                                    else:
                                        logger.warning('Kroki returned invalid PNG for fenced diagram')
                                        rendered = None
                                else:
                                    logger.warning('Kroki returned status %s for fenced mermaid render', resp.status_code)
                                    if resp.status_code == 400:
                                        error_msg = resp.text[:500] if hasattr(resp, 'text') and resp.text else 'Unknown error'
                                        logger.error('Kroki syntax error - fenced diagram code preview: %s', sanitized_block[:200])
                                        logger.error('Kroki error details: %s', error_msg)
                                    rendered = None
                            except Exception as e:
                                logger.exception('Kroki PNG rendering failed for fenced diagram: %s', e)
                                rendered = None

                            if rendered:
                                try:
                                    logger.info('Inserting locally rendered mermaid image into DOCX')
                                    if PIL_AVAILABLE:
                                        try:
                                            im = Image.open(BytesIO(rendered))
                                            im.verify()
                                        except Exception:
                                            logger.warning('Rendered mermaid bytes are not a valid image according to PIL; falling back to Kroki')
                                            rendered = None
                                    else:
                                        if not _is_png_bytes(rendered):
                                            logger.warning('Rendered mermaid bytes do not have a PNG header; falling back to Kroki')
                                            rendered = None

                                    if rendered:
                                        img_stream = BytesIO(rendered)
                                        img_stream.seek(0)
                                        try:
                                            doc.add_picture(img_stream, width=Inches(5))
                                        except Exception:
                                            logger.exception('doc.add_picture from BytesIO failed; trying temp file fallback')
                                            tmpname = None
                                            try:
                                                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tf:
                                                    tf.write(rendered)
                                                    tmpname = tf.name
                                                doc.add_picture(tmpname, width=Inches(5))
                                            finally:
                                                if tmpname:
                                                    try:
                                                        os.remove(tmpname)
                                                    except Exception:
                                                        pass

                                    if caption_text:
                                        cap_para = doc.add_paragraph()
                                        cap_run = cap_para.add_run(caption_text)
                                        cap_run.italic = True
                                        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                except Exception as e:
                                    logger.exception('Failed to insert locally rendered mermaid image into docx: %s', e)
                                    rendered = None

                            if not rendered:
                                try:
                                    kroki_url = 'https://kroki.io/mermaid/png'
                                    resp = httpx.post(kroki_url, content=sanitized_block.encode('utf-8'), headers={"Content-Type": "text/plain"}, timeout=30.0)
                                    if resp.status_code == 200:
                                        logger.info('Kroki returned image bytes for mermaid block')
                                        img_bytes = resp.content
                                        try:
                                            img_stream = BytesIO(img_bytes)
                                            doc.add_picture(img_stream, width=Inches(5))
                                            if caption_text:
                                                cap_para = doc.add_paragraph()
                                                cap_run = cap_para.add_run(caption_text)
                                                cap_run.italic = True
                                                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        except Exception as e:
                                            logger.warning('Failed to insert Kroki mermaid image into docx: %s', e)
                                            para = doc.add_paragraph(style='Normal')
                                            para.style.font.name = 'Consolas'
                                            para.style.font.size = Pt(10)
                                            para.text = block_text
                                    else:
                                        logger.warning('Kroki returned status %s for mermaid render', resp.status_code)
                                        para = doc.add_paragraph(style='Normal')
                                        para.style.font.name = 'Consolas'
                                        para.style.font.size = Pt(10)
                                        para.text = block_text
                                except Exception as e:
                                    logger.exception('Failed to call Kroki for mermaid rendering: %s', e)
                                    para = doc.add_paragraph(style='Normal')
                                    para.style.font.name = 'Consolas'
                                    para.style.font.size = Pt(10)
                                    para.text = block_text
                    else:
                        para = doc.add_paragraph(style='Normal')
                        para.style.font.name = 'Consolas'
                        para.style.font.size = Pt(10)
                        para.text = block_text
                code_block_lines = []
                in_code_block = False
                code_block_lang = None
            else:
                in_code_block = True
                code_block_lang = lang or None
            i += 1
            continue
        
        table_header_pattern = re.compile(r"^\s*\|.*\|\s*$")
        table_sep_pattern = re.compile(r"^\s*\|?\s*[:\-]+(?:\s*\|\s*[:\-]+)+\s*\|?\s*$")
        if i + 1 < len(lines) and table_header_pattern.match(line) and table_sep_pattern.match(lines[i + 1].strip()):
            header_line = line.strip().strip('|')
            header_cells = [h.strip() for h in re.split(r'\s*\|\s*', header_line)]
            try:
                current_table = _start_table(doc, header_cells)
                in_table = True
            except Exception as e:
                logger.warning("Failed to start table: %s", e)
                in_table = False
                current_table = None

            j = i + 2
            while j < len(lines):
                row_line = lines[j].strip()
                if not row_line or not row_line.startswith('|'):
                    break
                row_cells = [c.strip() for c in re.split(r'\s*\|\s*', row_line.strip().strip('|'))]
                try:
                    row = current_table.add_row()
                    for col_idx, cell_text in enumerate(row_cells[: len(header_cells) ]):
                        try:
                            cell = row.cells[col_idx]
                            cell.text = _clean_markdown_text(cell_text)
                        except Exception:
                            pass
                except Exception as e:
                    logger.debug("Failed to add table row: %s", e)
                j += 1

            if in_table and current_table:
                finalize_table(current_table)
            in_table = False
            current_table = None
            i = j
            continue

        if in_table and current_table:
            finalize_table(current_table)
            in_table = False
            current_table = None
        
        if stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            if stripped.startswith('- ') or stripped.startswith('* '):
                content = stripped[2:].strip()
                is_bullet = True
            else:
                content = re.sub(r'^\d+\.\s', '', stripped).strip()
                is_bullet = False
            
            if i + 1 < len(lines):
                next_line = lines[i + 1]
                next_stripped = next_line.strip()
                if (next_stripped and 
                    not next_stripped.startswith('- ') and 
                    not next_stripped.startswith('* ') and 
                    not re.match(r'^\d+\.\s', next_stripped)):
                    list_item_buffer.append(content)
                    list_item_buffer.append(next_stripped)
                    i += 2
                    while i < len(lines):
                        cont_line = lines[i]
                        cont_stripped = cont_line.strip()
                        if (not cont_stripped or 
                            cont_stripped.startswith('- ') or 
                            cont_stripped.startswith('* ') or 
                            re.match(r'^\d+\.\s', cont_stripped)):
                            break
                        list_item_buffer.append(cont_stripped)
                        i += 1
                    full_content = ' '.join(list_item_buffer)
                    list_item_buffer = []
                    full_content = _capitalize_sentence(full_content)
                    if is_bullet:
                        _add_bullet_paragraph(doc, full_content)
                    else:
                        para = doc.add_paragraph(style='List Number')
                        _add_formatted_text_to_paragraph(para, full_content)
                    continue
            
            content = _capitalize_sentence(content)
            if is_bullet:
                _add_bullet_paragraph(doc, content)
            else:
                para = doc.add_paragraph(style='List Number')
                _add_formatted_text_to_paragraph(para, content)
            i += 1
            continue
        
        _add_text_line(doc, stripped)
        i += 1
    
    if in_table and current_table:
        finalize_table(current_table)
    
    if in_code_block and code_block_lines:
        para = doc.add_paragraph(style='Normal')
        para.style.font.name = 'Consolas'
        para.style.font.size = Pt(10)
        para.text = '\n'.join(code_block_lines)

#function to strip markdown formatting from text
def _clean_markdown_text(text: str) -> str:
    if not text:
        return text

    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*\*([^*]*)\*\*', r'\1', text)
    text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'^---+$', '', text, flags=re.MULTILINE)
    text = re.sub(r'[ \t]+', ' ', text)
    text = text.strip()
    return text

#function to capitalize the first letter of a sentence if needed
def _capitalize_sentence(text: str) -> str:
    if not text:
        return text
    text = text.strip()
    if text and text[0].islower():
        return text[0].upper() + text[1:]
    return text

#function to add a heading with proper page break control
def _add_heading_with_break_control(doc, text: str, level: int):
    """Add a heading with page break control to prevent orphaned headings"""
    heading = doc.add_heading(text, level)
    try:
        heading.paragraph_format.keep_with_next = True
        heading.paragraph_format.widow_control = True
        heading.paragraph_format.keep_together = True
        if level <= 2:
            heading.paragraph_format.space_before = Pt(18)
        
    except Exception as e:
        logger.debug(f"Failed to set heading break properties: {e}")


#function to convert a single markdown line into appropriate DOCX elements
def _add_text_line(doc, line: str):
    stripped = line.strip()
    if not stripped:
        return
    
    if re.match(r'^---+$', stripped):
        return
    
    if stripped.startswith('#### '):
        header_text = _clean_markdown_text(stripped[5:])
        if header_text:
            _add_heading_with_break_control(doc, header_text, 4)
    elif stripped.startswith('### '):
        header_text = _clean_markdown_text(stripped[4:])
        if header_text:
            _add_heading_with_break_control(doc, header_text, 3)
    elif stripped.startswith('## '):
        header_text = _clean_markdown_text(stripped[3:])
        if header_text:
            _add_heading_with_break_control(doc, header_text, 2)
    elif stripped.startswith('# '):
        header_text = _clean_markdown_text(stripped[2:])
        if header_text:
            _add_heading_with_break_control(doc, header_text, 1)
    else:
        content = _clean_markdown_text(stripped)
        if content:
            if content and len(content) > 0 and content[0].islower() and content[0].isalpha():
                content = _capitalize_sentence(content)
            para = doc.add_paragraph()
            _add_formatted_text_to_paragraph(para, content)

#function to generate a DOCX file from responses and return bytes
def generate_rfp_docx(
    individual_responses: List[Dict[str, Any]],
    requirements_result: RequirementsResult,
    extraction_result: ExtractionResult,
    rfp_title: Optional[str] = None,
    output_path: Optional[Path] = None,
) -> bytes:
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    
    logger.info("Generating DOCX document: %d requirement responses", len(individual_responses))
    
    is_structured = len(individual_responses) == 1 and individual_responses[0].get('requirement_id') == 'STRUCTURED'
    toc_entries = []
    
    if is_structured:
        if (requirements_result.structure_detection and 
            requirements_result.structure_detection.detected_sections):
            toc_entries = [
                {'text': f"{idx + 1}. {section}", 'level': 1}
                for idx, section in enumerate(requirements_result.structure_detection.detected_sections)
            ]
        else:
            response_text = individual_responses[0].get('response', '')
            headings = _extract_headings_from_markdown(response_text)
            numbered_pattern = re.compile(r'^\d+\.\s+')
            toc_entries = [
                h for h in headings 
                if h['level'] <= 2 and numbered_pattern.match(h['text'])
            ]
    else:
        for idx, resp_data in enumerate(individual_responses, 1):
            req_id = resp_data.get('requirement_id', 'N/A')
            toc_entries.append({
                'text': f"Requirement {idx}: {req_id}",
                'level': 2,
            })
    
    doc = Document()
    
    setup_styles(doc)
    
    setup_page_formatting(doc, start_page_number=1)
    
    section = doc.sections[0]
    sect_pr = section._sectPr
    footer = section.footer
    for para in footer.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)
    
    project_root = Path(__file__).parent.parent.parent
    
    final_title = rfp_title or extraction_result.language.upper()
    add_modern_front_page(doc, final_title, project_root)
    
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.is_linked_to_previous = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    sect_pr = section._sectPr
    footer = section.footer
    for para in footer.paragraphs[:]:
        p = para._element
        p.getparent().remove(p)
    
    _add_heading_with_break_control(doc, "Table of Contents", 1)
    doc.add_paragraph()
    
    add_manual_toc(doc, toc_entries)
    
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    section.is_linked_to_previous = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    sect_pr = section._sectPr
    existing_pg_num = sect_pr.find(qn('w:pgNumType'))
    if existing_pg_num is not None:
        sect_pr.remove(existing_pg_num)
    pg_num_type = OxmlElement('w:pgNumType')
    pg_num_type.set(qn('w:start'), '1')
    pg_num_type.set(qn('w:fmt'), 'decimal')
    sect_pr.append(pg_num_type)
    
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    footer = section.footer
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
        p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)
    
    _add_heading_with_break_control(doc, "Company Overview", 1)
    overview_text = """At fusionAIx, we believe that the future of digital transformation lies in the seamless blend of low-code platforms and artificial intelligence. Our core team brings together decades of implementation experience, domain expertise, and a passion for innovation. We partner with enterprises to reimagine processes, accelerate application delivery, and unlock new levels of efficiency. We help businesses scale smarter, faster, and with greater impact.

With a collaborative spirit and a commitment to excellence, our team transforms complex challenges into intelligent, practical solutions. fusionAIx is not just about technology—it's about empowering people, industries, and enterprises to thrive in a digital-first world.

We are proud to be officially recognized as a Great Place To Work® Certified Company for 2025–26, reflecting our commitment to a culture built on trust, innovation, and people-first values.

fusionAIx delivers tailored solutions that blend AI and automation to drive measurable results across industries. We are a niche Pega partner with 20+ successful Pega Constellation implementations across the globe. As Constellation migration experts, we focus on pattern-based development with Constellation, enabling faster project go-lives than traditional implementation approaches.

Our proven capabilities span three core technology platforms: Pega Constellation, Microsoft Power Platform, and ServiceNow. Through these platforms, we provide comprehensive services including Low Code/No Code development, Digital Process Transformation, and AI & Data solutions.

To accelerate time-to-value, fusionAIx offers proprietary accelerators and solution components including fxAgentSDK, fxAIStudio, fxMockUpToView, and fxSmartDCO. These tools enable rapid development, intelligent automation, and streamlined project delivery.

We support clients across diverse industries including Insurance, Banking & Finance, Government & Public Sector, Automotive & Fleet Management, and Travel & Tourism, combining platform expertise with structured knowledge transfer to help customers build sustainable, future-ready capabilities."""
    
    for para_text in overview_text.split('\n\n'):
        if para_text.strip():
            doc.add_paragraph(para_text.strip())
    
    doc.add_page_break()
    
    is_structured = len(individual_responses) == 1 and individual_responses[0].get('requirement_id') == 'STRUCTURED'
    
    if is_structured:
        response_text = individual_responses[0].get('response', '')
        _parse_markdown_to_docx(doc, response_text)
    else:
        _add_heading_with_break_control(doc, "Solution Requirement Responses", 1)
        for idx, resp_data in enumerate(individual_responses, 1):
            _add_heading_with_break_control(doc, f"Requirement {idx}: {resp_data.get('requirement_id', 'N/A')}", 2)
            
            req_para = doc.add_paragraph()
            req_para.add_run("Requirement: ").bold = True
            req_text = resp_data.get('requirement_text', '')
            if req_text:
                req_para.add_run(_capitalize_sentence(req_text))
            
            _add_heading_with_break_control(doc, "Response", 3)
            _parse_markdown_to_docx(doc, resp_data.get('response', ''))
            
            if resp_data.get('quality'):
                quality = resp_data['quality']
                quality_para = doc.add_paragraph()
                quality_para.add_run(f"Quality Score: {quality.get('score', 0):.0f}/100 | ").bold = True
                quality_para.add_run(f"Completeness: {quality.get('completeness', 'unknown')} | ")
                quality_para.add_run(f"Relevance: {quality.get('relevance', 'unknown')}")
            
            doc.add_paragraph()
    
    if output_path:
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("DOCX saved to: %s", output_path.absolute())
        return output_path.read_bytes()
    else:
        buf = BytesIO()
        doc.save(buf)
        bytes_data = buf.getvalue()
        logger.info("DOCX generated: %d bytes (in memory)", len(bytes_data))
        return bytes_data

