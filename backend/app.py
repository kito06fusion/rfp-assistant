from __future__ import annotations

import logging
import time
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, Response
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
import httpx
import re
import asyncio

from backend.mermaid_renderer import render_mermaid_to_bytes

from backend.pipeline.text_extraction import extract_text_from_file
from backend.agents.preprocess_agent import run_preprocess_agent
from backend.agents.requirements_agent import run_requirements_agent
from backend.agents.build_query import build_query, build_query_for_single_requirement
from backend.agents.response_agent import run_response_agent
from backend.agents.structure_detection_agent import detect_structure
from backend.agents.structured_response_agent import run_structured_response_agent
from backend.agents.question_agent import (
    analyze_requirements_for_questions,
    analyze_build_query_for_questions,
    analyze_build_query_for_questions_legacy,
    infer_answered_questions_from_answer,
    get_next_critical_question,
    check_if_more_questions_needed,
)
from backend.agents.quality_agent import assess_response_quality
from backend.rag import RAGSystem
from backend.models import (
    ExtractionResult,
    RequirementsResult,
    StructureDetectionResult,
    Question,
    Answer,
    ConversationContext,
    BuildQuery,
    PreprocessResult,
)
from backend.knowledge_base import FusionAIxKnowledgeBase
from backend.knowledge_base.company_kb import CompanyKnowledgeBase
from backend.memory.mem0_client import (
    store_preprocess_result,
    store_requirements_result,
    store_build_query_result,
)


#function to setup rag system and load the fusionAIx knowledge base
def _setup_rag_and_kb(use_rag: bool) -> tuple[Optional[RAGSystem], FusionAIxKnowledgeBase]:
    rag_system = None
    if use_rag:
        try:
            project_root = Path(__file__).parent.parent
            docs_folder = project_root / "docs"
            index_path = project_root / "rag_index"
            query_cache_path = project_root / "rag_query_cache.pkl"
            rag_system = RAGSystem(
                docs_folder=str(docs_folder), 
                index_path=str(index_path),
                query_cache_path=str(query_cache_path),
            )
            try:
                rag_system.ensure_index_up_to_date()
                stats = rag_system.get_stats()
                logger.info(
                    "RAG index ready | built=%s, docs=%s, vectors=%s, dim=%s, model=%s",
                    stats.get("index_built"),
                    stats.get("num_documents"),
                    stats.get("num_vectors"),
                    stats.get("embedding_dimension"),
                    stats.get("embedding_model"),
                )
            except ValueError as build_err:
                logger.warning(
                    "Failed to build RAG index: %s. RAG system not available. "
                    "Make sure you have documents (PDF, DOCX, TXT) in the 'docs' folder. "
                    "Continuing without RAG.",
                    str(build_err)
                )
                rag_system = None
        except Exception as rag_exc:
            logger.warning("Failed to load/build RAG system: %s. Continuing without RAG.", rag_exc)
            rag_system = None
    
    knowledge_base = get_fusionaix_kb()
    logger.info(
        "Using fusionAIx knowledge base: %d capabilities, %d case studies, %d accelerators",
        len(knowledge_base.capabilities),
        len(knowledge_base.case_studies),
        len(knowledge_base.accelerators),
    )
    return rag_system, knowledge_base


#function to enrich a build query with compact RAG index previews
def _enrich_build_query_with_rag(
    build_query: BuildQuery,
    requirements_result: RequirementsResult,
    rag_contexts_by_req: Dict[str, str],
) -> BuildQuery:
    if not rag_contexts_by_req:
        return build_query

    logger.info(
        "Enriching build query with compact RAG index for %d requirement(s)",
        len(rag_contexts_by_req),
    )

    req_text_by_id: Dict[str, str] = {
        req.id: req.source_text for req in requirements_result.solution_requirements
    }

    base_text = build_query.query_text or ""

    rag_section_lines: List[str] = []
    rag_section_lines.append("")
    rag_section_lines.append("=" * 80)
    rag_section_lines.append("RAG-SUPPORTED REQUIREMENTS (compact index only – full evidence kept outside build_query)")
    rag_section_lines.append("=" * 80)
    rag_section_lines.append("")

    for req_id, rag_ctx in rag_contexts_by_req.items():
        req_text = req_text_by_id.get(req_id, "")
        rag_section_lines.append(f"Requirement ID: {req_id}")
        if req_text:
            rag_section_lines.append(f"Requirement: {req_text}")
        first_line = rag_ctx.strip().splitlines()[0] if rag_ctx.strip() else ""
        preview = (first_line[:140] + "...") if len(first_line) > 140 else first_line
        rag_section_lines.append(
            f"RAG: evidence available in RAG system "
            f"(preview: {preview or 'no preview available'})"
        )
        rag_section_lines.append("--- END RAG INDEX ITEM ---")
        rag_section_lines.append("")

    enriched_text = base_text.rstrip() + "\n" + "\n".join(rag_section_lines)
    build_query.query_text = enriched_text
    return build_query


#function to validate extraction and requirements before generating responses
def validate_before_generation(
    extraction_result: ExtractionResult,
    requirements_result: RequirementsResult,
) -> List[str]:
    errors = []
    
    if not extraction_result.language:
        errors.append("Extraction result missing language")
    if not isinstance(extraction_result.language, str) or len(extraction_result.language) < 2:
        errors.append(f"Invalid language code: {extraction_result.language}")
    
    if not requirements_result.solution_requirements:
        errors.append("No solution requirements found")
        return errors
    
    for idx, req in enumerate(requirements_result.solution_requirements, 1):
        if not req.id or not req.id.strip():
            errors.append(f"Solution requirement {idx} missing ID")
        if not req.source_text or not req.source_text.strip():
            errors.append(f"Solution requirement {idx} ({req.id}) missing source_text")
        if not req.category or not req.category.strip():
            errors.append(f"Solution requirement {idx} ({req.id}) missing category")
        if req.source_text and len(req.source_text) < 10:
            errors.append(f"Solution requirement {idx} ({req.id}) source_text too short (likely incomplete)")
    
    if not requirements_result.response_structure_requirements:
        logger.warning("No response structure requirements found - responses may lack structure guidance")
    else:
        for idx, req in enumerate(requirements_result.response_structure_requirements, 1):
            if not req.id or not req.id.strip():
                errors.append(f"Response structure requirement {idx} missing ID")
            if not req.source_text or not req.source_text.strip():
                errors.append(f"Response structure requirement {idx} ({req.id}) missing source_text")
    
    if requirements_result.solution_requirements:
        try:
            test_build_query = build_query_for_single_requirement(
                extraction_result=extraction_result,
                single_requirement=requirements_result.solution_requirements[0],
                all_response_structure_requirements=requirements_result.response_structure_requirements,
            )
            if not test_build_query.query_text or len(test_build_query.query_text) < 100:
                errors.append("Build query test failed - generated query text too short")
            if not test_build_query.solution_requirements_summary:
                errors.append("Build query test failed - missing solution requirements summary")
        except Exception as e:
            errors.append(f"Build query test failed: {str(e)}")
    
    solution_ids = [req.id for req in requirements_result.solution_requirements]
    if len(solution_ids) != len(set(solution_ids)):
        duplicates = [id for id in solution_ids if solution_ids.count(id) > 1]
        errors.append(f"Duplicate solution requirement IDs found: {set(duplicates)}")
    
    response_ids = [req.id for req in requirements_result.response_structure_requirements]
    if len(response_ids) != len(set(response_ids)):
        duplicates = [id for id in response_ids if response_ids.count(id) > 1]
        errors.append(f"Duplicate response structure requirement IDs found: {set(duplicates)}")
    
    return errors


LOG_FORMAT = "%(asctime)s [%(levelname)s] %(name)s: %(message)s"

logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)
logger = logging.getLogger(__name__)


app = FastAPI(title="RFP Assistant Backend")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

project_root = Path(__file__).parent.parent
frontend_dist = project_root / "frontend" / "dist"
frontend_src = project_root / "frontend" / "src"

if frontend_dist.exists() and (frontend_dist / "index.html").exists():
    if (frontend_dist / "assets").exists():
        app.mount("/assets", StaticFiles(directory=str(frontend_dist / "assets")), name="assets")
    logger.info("Serving frontend from dist directory (production build)")
elif frontend_src.exists():
    app.mount("/src", StaticFiles(directory=str(frontend_src)), name="src")
    if (frontend_src.parent / "public").exists():
        app.mount("/public", StaticFiles(directory=str(frontend_src.parent / "public")), name="public")
    logger.info("Serving frontend from src directory (development mode - build frontend for production)")

_fusionaix_kb: FusionAIxKnowledgeBase | None = None

#function to return a cached fusionAIx knowledge base instance
def get_fusionaix_kb() -> FusionAIxKnowledgeBase:
    global _fusionaix_kb
    if _fusionaix_kb is None:
        logger.info("Initializing fusionAIx knowledge base")
        _fusionaix_kb = FusionAIxKnowledgeBase()
        logger.info(
            "Knowledge base loaded: %d capabilities, %d case studies, %d accelerators",
            len(_fusionaix_kb.capabilities),
            len(_fusionaix_kb.case_studies),
            len(_fusionaix_kb.accelerators),
        )
    return _fusionaix_kb


#function to serve the frontend index.html at root
@app.get("/", response_class=HTMLResponse)
async def index() -> HTMLResponse:
    project_root = Path(__file__).resolve().parent.parent
    index_path = project_root / "frontend" / "dist" / "index.html"
    if not index_path.exists():
        index_path = project_root / "frontend" / "index.html"
    if not index_path.exists():
        raise HTTPException(status_code=404, detail="Frontend not found. Please build the frontend first: npm run build")
    return HTMLResponse(index_path.read_text(encoding="utf-8"))


SUPPORTED_FILE_TYPES = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".txt"}


#function to process uploaded RFP files and extract combined text
@app.post("/process-rfp")
async def process_rfp(files: List[UploadFile] = File(...)) -> Dict[str, Any]:
    request_id = str(uuid.uuid4())
    file_names = [f.filename for f in files]
    logger.info("REQUEST %s: /process-rfp files=%s", request_id, file_names)

    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded.")

    for file in files:
        suffix = Path(file.filename).suffix.lower()
        if suffix not in SUPPORTED_FILE_TYPES:
            logger.warning("REQUEST %s: unsupported file type %s in file %s", request_id, suffix, file.filename)
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Unsupported file type '{suffix}' for file '{file.filename}'. "
                    f"Please upload: PDF, DOCX, DOC, XLSX, XLS, or TXT."
                ),
            )

    t0 = time.time()
    all_text_parts: List[str] = []
    temp_paths: List[Path] = []

    for file_index, file in enumerate(files, 1):
        temp_path = Path("/tmp") / f"{request_id}_{file_index}_{file.filename}"
        temp_paths.append(temp_path)
        bytes_written = 0

        try:
            with open(temp_path, "wb") as f:
                while True:
                    chunk = await file.read(2 * 1024 * 1024)
                    if not chunk:
                        break
                    f.write(chunk)
                    bytes_written += len(chunk)
            logger.info(
                "REQUEST %s: uploaded file %d/%d: %s (%d bytes)",
                request_id, file_index, len(files), file.filename, bytes_written,
            )
        except Exception as e:
            logger.exception("REQUEST %s: failed to write file to disk: %s", request_id, e)
            for tp in temp_paths:
                try:
                    tp.unlink(missing_ok=True)
                except Exception:
                    pass
            raise HTTPException(status_code=500, detail="Failed to process uploaded file.") from e

        try:
            file_text = extract_text_from_file(temp_path)
            logger.info(
                "REQUEST %s: extracted %d chars from file %d/%d (%s)",
                request_id, len(file_text), file_index, len(files), file.filename,
            )
            if file_text.strip():
                if len(files) > 1:
                    all_text_parts.append(f"\n{'='*60}\nFILE: {file.filename}\n{'='*60}\n\n{file_text}")
                else:
                    all_text_parts.append(file_text)
        except Exception as e:
            logger.exception("REQUEST %s: failed to extract text from %s: %s", request_id, file.filename, e)

    for temp_path in temp_paths:
        try:
            temp_path.unlink(missing_ok=True)
        except Exception:
            pass

    text = "\n\n".join(all_text_parts)
    logger.info("REQUEST %s: combined text from %d file(s): %d chars", request_id, len(files), len(text))

    if not text.strip():
        logger.warning("REQUEST %s: no text extracted from files", request_id)
        raise HTTPException(status_code=400, detail="No text could be extracted from the uploaded files.")

    elapsed = time.time() - t0
    logger.info("REQUEST %s: OCR extraction completed in %.2fs", request_id, elapsed)
    
    response: Dict[str, Any] = {
        "preprocess": None,
        "requirements": None,
        "ocr_source_text": text,
    }
    return response


class PreprocessRequest(BaseModel):
    ocr_text: str


class RenderRequest(BaseModel):
    diagram: str
    format: str = 'png'


#function to run the preprocess agent on OCR text
@app.post("/run-preprocess")
async def run_preprocess(req: PreprocessRequest) -> Dict[str, Any]:
    request_id = str(uuid.uuid4())
    logger.info("REQUEST %s: /run-preprocess called with %d chars", request_id, len(req.ocr_text))
    
    if not req.ocr_text.strip():
        raise HTTPException(status_code=400, detail="OCR text is empty.")
    
    t0 = time.time()
    try:
        preprocess_res = run_preprocess_agent(req.ocr_text)
    except Exception as exc:
        elapsed = time.time() - t0
        logger.exception(
            "REQUEST %s: preprocess agent failed after %.2fs: %s",
            request_id,
            elapsed,
            exc,
        )
        raise HTTPException(
            status_code=500,
            detail=f"Preprocess agent failed for request {request_id}. Check server logs.",
        ) from exc
    
    elapsed = time.time() - t0
    logger.info(
        "REQUEST %s: preprocess agent finished (cleaned_chars=%d, removed_chars=%d, comparison_agreement=%s)",
        request_id,
        len(preprocess_res.cleaned_text or ""),
        len(preprocess_res.removed_text or ""),
        preprocess_res.comparison_agreement,
    )
    logger.info("REQUEST %s: preprocess completed in %.2fs", request_id, elapsed)
    
    response: Dict[str, Any] = preprocess_res.to_dict()
    response["ocr_text"] = req.ocr_text

    try:
        store_preprocess_result(req.ocr_text, preprocess_res.to_dict())
    except Exception as exc:
        logger.warning("Mem0 preprocess storage failed: %s", exc)

    return response


#function to render mermaid diagrams (via local renderer or Kroki fallback)
@app.post('/render/mermaid')
async def render_mermaid(req: RenderRequest):
    diagram = (req.diagram or "").strip()
    fmt = (req.format or 'png').lower()
    if not diagram:
        raise HTTPException(status_code=400, detail='No diagram provided')
    if fmt not in ('png', 'svg'):
        raise HTTPException(status_code=400, detail='Invalid format: choose png or svg')

    try:
        logger.info('Received raw diagram (escaped newlines): %s', diagram.replace('\n','\\n')[:1000])
    except Exception:
        logger.debug('Failed to log raw diagram')

    caption = None
    m = re.search(r'^(?P<diag>.*?)(?:\n+Caption:\s*(?P<cap>.*))?$', diagram, flags=re.I | re.S)
    if m:
        diagram_text = (m.group('diag') or '').rstrip()
        caption = (m.group('cap') or '').strip() or None
        if caption:
            logger.info('Detected and stripped caption from diagram for rendering: %s', caption)
    try:
        preview = diagram_text.replace('\n', '\\n')[:1000]
        logger.info('Diagram text preview (escaped newlines): %s', preview)
        logger.info('Diagram text line count: %d', len(diagram_text.splitlines()))
    except Exception:
        logger.debug('Failed to produce diagram preview for logs')
    
    try:
        loop = asyncio.get_running_loop()
        logger.info('Attempting MCP mermaid rendering for diagram (fmt=%s)', fmt)
        data = await loop.run_in_executor(None, lambda: render_mermaid_to_bytes(diagram_text, fmt))
        media_type = 'image/png' if fmt == 'png' else 'image/svg+xml'
        logger.info('MCP mermaid rendering succeeded; returning image bytes')
        headers = {"X-Diagram-Renderer": "mcp-mermaid"}
        if caption:
            headers["X-Diagram-Caption"] = caption
        return Response(content=data, media_type=media_type, headers=headers)
    except Exception as e:
        logger.exception('MCP mermaid rendering failed: %s', e)

    kroki_url = f'https://kroki.io/mermaid/{fmt}'

    try:
        logger.info('Requesting Kroki rendering as fallback (fmt=%s)', fmt)
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(kroki_url, content=diagram_text.encode('utf-8'), headers={"Content-Type": "text/plain"})
    except Exception as e:
        logger.exception('Kroki request failed: %s', e)
        raise HTTPException(status_code=500, detail='Failed to contact Kroki rendering service')

    if resp.status_code != 200:
        logger.error('Kroki returned status %s: %s', resp.status_code, resp.text[:200])
        raise HTTPException(status_code=502, detail='Kroki rendering failed')

    media_type = 'image/png' if fmt == 'png' else 'image/svg+xml'
    logger.info('Kroki rendering succeeded; returning image bytes')
    headers = {"X-Diagram-Renderer": "kroki"}
    if caption:
        headers["X-Diagram-Caption"] = caption
    return Response(content=resp.content, media_type=media_type, headers=headers)


class RequirementsRequest(BaseModel):
    essential_text: str


class UpdateRequirementsRequest(BaseModel):
    requirements: Dict[str, Any]

#function to run requirements agent and detect structure of responses
@app.post("/run-requirements")
async def run_requirements(req: RequirementsRequest) -> Dict[str, Any]:
    logger.info(
        "Requirements endpoint called (essential_chars=%d)", len(req.essential_text)
    )
    try:
        result = run_requirements_agent(essential_text=req.essential_text)
        
        logger.info("Running structure detection on %d response structure requirements", 
                   len(result.response_structure_requirements))
        structure_detection_dict = detect_structure(result.response_structure_requirements)
        structure_detection = StructureDetectionResult(**structure_detection_dict)
        result.structure_detection = structure_detection
        
        logger.info(
            "Structure detection completed: explicit=%s, type=%s, sections=%d, confidence=%.2f",
            structure_detection.has_explicit_structure,
            structure_detection.structure_type,
            len(structure_detection.detected_sections),
            structure_detection.confidence,
        )
    except Exception as exc:
        logger.exception("Requirements agent failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="Requirements agent failed. Check server logs.",
        ) from exc
    logger.info(
        "Requirements agent completed (solution=%d, response_structure=%d)",
        len(result.solution_requirements),
        len(result.response_structure_requirements),
    )

    try:
        store_requirements_result(req.essential_text, result.to_dict())
    except Exception as exc:
        logger.warning("Mem0 requirements storage failed: %s", exc)

    return result.to_dict()


#function to update and validate requirements payload
@app.post("/update-requirements")
async def update_requirements(req: UpdateRequirementsRequest) -> Dict[str, Any]:
    logger.info("Update requirements endpoint called")
    try:
        result = RequirementsResult(**req.requirements)
        return result.to_dict()
    except Exception as exc:
        logger.exception("Update requirements failed: %s", exc)
        raise HTTPException(
            status_code=400,
            detail=f"Invalid requirements payload: {str(exc)}",
        ) from exc

class BuildQueryRequest(BaseModel):
    preprocess: Dict[str, Any]
    requirements: Dict[str, Any]


#function to extract a short title from key requirements summary
def _extract_title_from_key_requirements(key_requirements_summary: str) -> Optional[str]:
    if not key_requirements_summary:
        return None
    
    lines = key_requirements_summary.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        for marker in ['-', '•', '*', '·']:
            if line.startswith(marker):
                title = line[1:].strip()
                if title:
                    return title
        return line
    
    return None


#function to convert a PreprocessResult into an ExtractionResult
def _extraction_from_preprocess(pre: PreprocessResult) -> ExtractionResult:
    return ExtractionResult(
        translated_text="",
        language=pre.language,
        key_requirements_summary=pre.key_requirements_summary,
        raw_structured={},
    )


#function to build combined text response from individual responses (for logging)
def _build_combined_response_text(
    individual_responses: List[Dict[str, Any]],
    requirements_result: RequirementsResult,
) -> str:
    combined_parts = []
    combined_parts.append("=" * 80)
    combined_parts.append("RFP RESPONSE DOCUMENT")
    combined_parts.append("=" * 80)
    combined_parts.append("")
    
    if requirements_result.response_structure_requirements:
        combined_parts.append("RESPONSE STRUCTURE REQUIREMENTS")
        combined_parts.append("-" * 80)
        for resp_req in requirements_result.response_structure_requirements:
            combined_parts.append(resp_req.source_text)
            combined_parts.append("")
        combined_parts.append("")
    
    combined_parts.append("SOLUTION REQUIREMENT RESPONSES")
    combined_parts.append("=" * 80)
    combined_parts.append("")
    
    for idx, resp_data in enumerate(individual_responses, 1):
        combined_parts.append(f"Requirement {idx}: {resp_data['requirement_id']}")
        combined_parts.append("-" * 80)
        combined_parts.append(f"Requirement: {resp_data['requirement_text']}")
        combined_parts.append("")
        combined_parts.append("Response:")
        combined_parts.append("-" * 40)
        combined_parts.append(resp_data['response'])
        combined_parts.append("")
        combined_parts.append("")
    
    return "\n".join(combined_parts)


#function to extract key phrase from requirement text
def _extract_key_phrase(source_text: str, max_words: int = 10) -> str:
    words = source_text.split()
    return " ".join(words[:max_words]) + ("..." if len(words) > max_words else "")


#function to create an error response dict for a failed requirement
def _create_error_response(
    requirement_id: str,
    requirement_text: str,
    key_phrase: str,
    error: Exception,
) -> Dict[str, Any]:
    return {
        "requirement_id": requirement_id,
        "requirement_text": requirement_text,
        "key_phrase": key_phrase,
        "response": f"[ERROR: Failed to generate response for this requirement: {str(error)}]",
        "notes": f"Error: {str(error)}",
        "quality": {
            "score": 0.0,
            "completeness": "incomplete",
            "relevance": "low",
            "issues": [f"Generation failed: {str(error)}"],
            "suggestions": ["Fix the error and regenerate"],
        },
    }


#function to log progress for requirement processing
def _log_requirement_progress(
    idx: int,
    total: int,
    successful: int,
    failed: int,
    requirement_id: str,
) -> None:
    logger.info(
        "[%d/%d] Progress: %d successful, %d failed, %d remaining",
        idx,
        total,
        successful,
        failed,
        total - idx,
    )


#function to log final generation summary
def _log_generation_summary(
    total_requirements: int,
    successful_responses: int,
    failed_responses: int,
    total_elapsed: float,
    combined_response_length: int,
) -> None:
    logger.info("=" * 80)
    logger.info("Response generation completed:")
    logger.info("  Total requirements: %d", total_requirements)
    logger.info("  Successful responses: %d", successful_responses)
    logger.info("  Failed responses: %d", failed_responses)
    logger.info(
        "  Success rate: %.1f%%",
        (successful_responses / total_requirements * 100) if total_requirements > 0 else 0.0,
    )
    logger.info("  Total time: %.2f seconds", total_elapsed)
    logger.info(
        "  Average time per requirement: %.2f seconds",
        total_elapsed / total_requirements if total_requirements > 0 else 0.0,
    )
    logger.info("  Combined response length: %d characters", combined_response_length)
    logger.info("=" * 80)


#function to process a single requirement and generate response
def _process_single_requirement(
    solution_req: Any,
    idx: int,
    total_requirements: int,
    extraction_result: ExtractionResult,
    requirements_result: RequirementsResult,
    knowledge_base: Any,
    qa_context: str,
) -> Dict[str, Any]:
    req_start_time = time.time()
    key_phrase = _extract_key_phrase(solution_req.source_text)
    
    logger.info(
        "[%d/%d] Processing requirement: %s",
        idx,
        total_requirements,
        solution_req.id,
    )
    logger.debug(
        "[%d/%d] Requirement text: %s",
        idx,
        total_requirements,
        solution_req.source_text[:100] + "..."
        if len(solution_req.source_text) > 100
        else solution_req.source_text,
    )
    
    build_query_obj = build_query_for_single_requirement(
        extraction_result=extraction_result,
        single_requirement=solution_req,
        all_response_structure_requirements=requirements_result.response_structure_requirements,
    )
    build_query_obj.confirmed = True
    
    try:
        result = run_response_agent(
            build_query=build_query_obj,
            knowledge_base=knowledge_base,
            qa_context=qa_context,
        )
        
        quality_assessment = assess_response_quality(solution_req, result.response_text)
        req_elapsed = time.time() - req_start_time
        
        response_dict = {
            "requirement_id": solution_req.id,
            "requirement_text": solution_req.source_text,
            "key_phrase": key_phrase,
            "response": result.response_text,
            "notes": result.notes,
            "quality": quality_assessment,
        }
        
        logger.info(
            "[%d/%d] ✓ SUCCESS: Generated response for requirement %s (length=%d chars, time=%.2fs)",
            idx,
            total_requirements,
            solution_req.id,
            len(result.response_text),
            req_elapsed,
        )
        
        return {"response": response_dict, "success": True, "elapsed": req_elapsed}
        
    except Exception as req_exc:
        req_elapsed = time.time() - req_start_time
        logger.error(
            "[%d/%d] ✗ FAILED: Requirement %s failed after %.2fs: %s",
            idx,
            total_requirements,
            solution_req.id,
            req_elapsed,
            req_exc,
        )
        logger.exception(
            "[%d/%d] Full error traceback for requirement %s:",
            idx,
            total_requirements,
            solution_req.id,
        )
        
        error_response = _create_error_response(
            requirement_id=solution_req.id,
            requirement_text=solution_req.source_text,
            key_phrase=key_phrase,
            error=req_exc,
        )
        
        return {"response": error_response, "success": False, "elapsed": req_elapsed}


#function to build a query from preprocess and requirements
@app.post("/build-query")
async def build_query_endpoint(req: BuildQueryRequest) -> Dict[str, Any]:
    logger.info("Build query endpoint called")
    try:
        preprocess_result = PreprocessResult(**req.preprocess)
        extraction_result = _extraction_from_preprocess(preprocess_result)
        requirements_result = RequirementsResult(**req.requirements)
        query = build_query(extraction_result, requirements_result)
        payload = query.model_dump()

        try:
            store_build_query_result(req.preprocess.get("cleaned_text") or "", payload)
        except Exception as exc:
            logger.warning("Mem0 build query storage failed: %s", exc)

        return payload
    except Exception as exc:
        logger.exception("Build query failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="Build query failed. Check server logs.",
        ) from exc


class GenerateResponseRequest(BaseModel):
    preprocess: Dict[str, Any]
    requirements: Dict[str, Any]
    use_rag: bool = True
    num_retrieval_chunks: int = 5
    session_id: Optional[str] = None


#function to generate responses (structured or per-requirement) for requirements
@app.post("/generate-response")
async def generate_response_endpoint(req: GenerateResponseRequest) -> Dict[str, Any]:
    logger.info("Generate response endpoint called (use_rag=%s)", req.use_rag)
    try:
        preprocess_result = PreprocessResult(**req.preprocess)
        extraction_result = _extraction_from_preprocess(preprocess_result)
        requirements_result = RequirementsResult(**req.requirements)
        
        if not requirements_result.solution_requirements:
            raise HTTPException(
                status_code=400,
                detail="No solution requirements found. Cannot generate response.",
            )
        
        if requirements_result.structure_detection is None:
            logger.info("Structure detection not found in requirements, running now...")
            structure_detection_dict = detect_structure(requirements_result.response_structure_requirements)
            requirements_result.structure_detection = StructureDetectionResult(**structure_detection_dict)
        
        structure_detection = requirements_result.structure_detection
        
        if structure_detection.has_explicit_structure and structure_detection.confidence >= 0.6:
            logger.info(
                "=" * 80
            )
            logger.info(
                "EXPLICIT STRUCTURE DETECTED - Using structured response generation"
            )
            logger.info(
                "Structure: %s (%d sections, confidence=%.2f)",
                structure_detection.structure_type,
                len(structure_detection.detected_sections),
                structure_detection.confidence,
            )
            logger.info(
                "=" * 80
            )
            return await _generate_structured_response(
                extraction_result=extraction_result,
                requirements_result=requirements_result,
                structure_detection=structure_detection,
                use_rag=req.use_rag,
                num_retrieval_chunks=req.num_retrieval_chunks,
                session_id=req.session_id,
            )
        else:
            logger.info(
                "=" * 80
            )
            logger.info(
                "NO EXPLICIT STRUCTURE - Using per-requirement response generation"
            )
            logger.info(
                "Structure type: %s, confidence: %.2f",
                structure_detection.structure_type if structure_detection else "none",
                structure_detection.confidence if structure_detection else 0.0,
            )
            logger.info(
                "=" * 80
            )
            return await _generate_per_requirement_response(
                extraction_result=extraction_result,
                requirements_result=requirements_result,
                use_rag=req.use_rag,
                num_retrieval_chunks=req.num_retrieval_chunks,
                session_id=req.session_id,
            )
    except HTTPException:
        raise
    except KeyboardInterrupt:
        raise
    except Exception as exc:
        logger.exception("Response generation failed completely: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Response generation failed. Check server logs. Error: {str(exc)}",
        ) from exc


#function to generate a single structured response and return a DOCX
async def _generate_structured_response(
    extraction_result: ExtractionResult,
    requirements_result: RequirementsResult,
    structure_detection: StructureDetectionResult,
    use_rag: bool,
    num_retrieval_chunks: int,
    session_id: Optional[str] = None,
) -> Response:
    rag_system, knowledge_base = _setup_rag_and_kb(use_rag)
    
    logger.info("=" * 80)
    logger.info("Running pre-flight validation before structured response generation...")
    validation_errors = validate_before_generation(
        extraction_result=extraction_result,
        requirements_result=requirements_result,
    )
    if validation_errors:
        error_msg = "Pre-flight validation failed. Please fix the following issues:\n" + "\n".join(f"  - {err}" for err in validation_errors)
        logger.error(error_msg)
        raise HTTPException(
            status_code=400,
            detail=error_msg,
        )
    logger.info("✓ Pre-flight validation passed - all checks OK")
    logger.info("=" * 80)
    
    logger.info("Generating structured response...")
    start_time = time.time()
    
    try:
        qa_context = ""
        if session_id and session_id in _conversation_sessions:
            context = _conversation_sessions[session_id]
            qa_context = context.get_qa_context()
            logger.info("Including Q&A context from session %s (%d answers)", session_id, len(context.answers))
        
        result = run_structured_response_agent(
            requirements_result=requirements_result,
            structure_detection=structure_detection,
            rag_system=rag_system,
            num_retrieval_chunks=num_retrieval_chunks,
            knowledge_base=knowledge_base,
            qa_context=qa_context,
        )
        
        individual_responses = [{
            "requirement_id": "STRUCTURED",
            "requirement_text": f"Complete structured response following: {', '.join(structure_detection.detected_sections)}",
            "key_phrase": structure_detection.detected_sections[0] if structure_detection.detected_sections else "Structured Response",
            "response": result.response_text,
            "notes": result.notes,
        }]
        
        total_elapsed = time.time() - start_time
        logger.info("Structured response generated in %.2f seconds (length=%d chars)", total_elapsed, len(result.response_text))
        
        try:
            from backend.document_formatter import generate_rfp_docx
        except ImportError as import_err:
            logger.error("DOCX generation not available: %s", import_err)
            raise HTTPException(
                status_code=500,
                detail="DOCX generation not available. Install python-docx to enable Word export.",
            ) from import_err
        
        rfp_title = _extract_title_from_key_requirements(extraction_result.key_requirements_summary)
        
        #generate DOCX in memory without saving to disk
        docx_bytes = generate_rfp_docx(
            individual_responses=individual_responses,
            requirements_result=requirements_result,
            extraction_result=extraction_result,
            rfp_title=rfp_title,
            output_path=None,
        )
        
        logger.info("DOCX generated successfully: %d bytes (in memory, not saved)", len(docx_bytes))
        
        timestamp = int(time.time())
        docx_filename = f"rfp_response_structured_{extraction_result.language}_{timestamp}.docx"
        
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'inline; filename="{docx_filename}"',
                "Content-Length": str(len(docx_bytes)),
            }
        )
    except Exception as exc:
        logger.exception("Structured response generation failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Structured response generation failed: {str(exc)}",
        ) from exc


#function to assemble individual responses into a DOCX response file
def _generate_docx_response(
    individual_responses: List[Dict[str, Any]],
    extraction_result: ExtractionResult,
    requirements_result: RequirementsResult,
) -> Response:
    try:
        from backend.document_formatter import generate_rfp_docx
    except ImportError as import_err:
        raise ImportError("DOCX generation not available. Install python-docx.") from import_err
    
    rfp_title = _extract_title_from_key_requirements(extraction_result.key_requirements_summary)
    
    logger.info("Generating DOCX with %d individual responses", len(individual_responses))
    docx_start_time = time.time()
    
    #generate DOCX in memory without saving to disk
    docx_bytes = generate_rfp_docx(
        individual_responses=individual_responses,
        requirements_result=requirements_result,
        extraction_result=extraction_result,
        rfp_title=rfp_title,
        output_path=None,
    )
    
    docx_elapsed = time.time() - docx_start_time
    timestamp = int(time.time())
    docx_filename = f"rfp_response_{extraction_result.language}_{timestamp}.docx"
    
    logger.info(
        "DOCX generation completed successfully: %d bytes in %.2f seconds (in memory, not saved)",
        len(docx_bytes),
        docx_elapsed,
    )
    
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'inline; filename="{docx_filename}"',
            "Content-Length": str(len(docx_bytes)),
        }
    )


#function to generate responses per solution requirement and return DOCX
async def _generate_per_requirement_response(
    extraction_result: ExtractionResult,
    requirements_result: RequirementsResult,
    use_rag: bool,
    num_retrieval_chunks: int,
    session_id: Optional[str] = None,
) -> Response:
    rag_system, knowledge_base = _setup_rag_and_kb(use_rag)
    
    qa_context = ""
    if session_id and session_id in _conversation_sessions:
        context = _conversation_sessions[session_id]
        qa_context = context.get_qa_context()
        logger.info("Including Q&A context from session %s (%d answers)", session_id, len(context.answers))
    
    logger.info("=" * 80)
    logger.info("Running pre-flight validation before response generation...")
    validation_errors = validate_before_generation(
        extraction_result=extraction_result,
        requirements_result=requirements_result,
    )
    if validation_errors:
        error_msg = "Pre-flight validation failed. Please fix the following issues:\n" + "\n".join(f"  - {err}" for err in validation_errors)
        logger.error(error_msg)
        raise HTTPException(
            status_code=400,
            detail=error_msg,
        )
    logger.info("✓ Pre-flight validation passed - all checks OK")
    logger.info("=" * 80)
    
    total_requirements = len(requirements_result.solution_requirements)
    logger.info("=" * 80)
    logger.info("Starting response generation for %d solution requirement(s)", total_requirements)
    logger.info("=" * 80)
    
    individual_responses = []
    successful_responses = 0
    failed_responses = 0
    start_time = time.time()
    partial_completion = False
    
    try:
        for idx, solution_req in enumerate(requirements_result.solution_requirements, 1):
            result = _process_single_requirement(
                solution_req=solution_req,
                idx=idx,
                total_requirements=total_requirements,
                extraction_result=extraction_result,
                requirements_result=requirements_result,
                knowledge_base=knowledge_base,
                qa_context=qa_context,
            )
            
            individual_responses.append(result["response"])
            
            if result["success"]:
                successful_responses += 1
            else:
                failed_responses += 1
            
            _log_requirement_progress(
                idx=idx,
                total=total_requirements,
                successful=successful_responses,
                failed=failed_responses,
                requirement_id=solution_req.id,
            )
            
    except KeyboardInterrupt:
        logger.warning("Response generation interrupted by user")
        partial_completion = True
        raise
    except Exception as catastrophic_error:
        logger.error("Catastrophic error during response generation: %s", catastrophic_error)
        logger.exception("Full traceback:")
        partial_completion = True
        if not individual_responses:
            raise HTTPException(
                status_code=500,
                detail="Response generation failed before any responses could be generated. Check server logs.",
            ) from catastrophic_error
    
    if not individual_responses:
        raise HTTPException(
            status_code=500,
            detail="Response generation failed before any responses could be generated. Check server logs.",
        )
    
    if partial_completion:
        logger.warning(
            "Response generation completed partially: %d/%d requirements processed",
            len(individual_responses),
            total_requirements,
        )
    
    total_elapsed = time.time() - start_time
    combined_response_text = _build_combined_response_text(individual_responses, requirements_result)
    _log_generation_summary(
        total_requirements=total_requirements,
        successful_responses=successful_responses,
        failed_responses=failed_responses,
        total_elapsed=total_elapsed,
        combined_response_length=len(combined_response_text),
    )
    
    logger.info("Generating DOCX document...")
    try:
        return _generate_docx_response(individual_responses, extraction_result, requirements_result)
    except ImportError as import_exc:
        logger.error("DOCX generation not available: %s", import_exc)
        raise HTTPException(
            status_code=500,
            detail="DOCX generation not available. python-docx dependency is missing. Check server logs.",
        ) from import_exc
    except Exception as docx_exc:
        logger.exception("DOCX generation failed: %s", docx_exc)
        raise HTTPException(
            status_code=500,
            detail=f"DOCX generation failed: {str(docx_exc)}",
        ) from docx_exc


_conversation_sessions: Dict[str, ConversationContext] = {}
_company_kb_instance: Optional[CompanyKnowledgeBase] = None

#function to return a cached company knowledge base instance
def get_company_kb() -> CompanyKnowledgeBase:
    global _company_kb_instance
    if _company_kb_instance is None:
        logger.info("Initializing company knowledge base")
        _company_kb_instance = CompanyKnowledgeBase()
    return _company_kb_instance


class GenerateQuestionsRequest(BaseModel):
    requirements: Optional[Dict[str, Any]] = None
    build_query: Optional[Dict[str, Any]] = None


class EnrichBuildQueryRequest(BaseModel):
    build_query: Dict[str, Any]
    session_id: Optional[str] = None


#function to generate clarification questions from build_query or requirements
@app.post("/generate-questions")
async def generate_questions_endpoint(req: GenerateQuestionsRequest) -> Dict[str, Any]:
    logger.info("Generate questions endpoint called")
    try:
        company_kb = get_company_kb()
        rag_system, _ = _setup_rag_and_kb(use_rag=True)
        
        if req.build_query:
            logger.info("Analyzing build query for questions")
            build_query_obj = BuildQuery(**req.build_query)
            
            if req.requirements:
                requirements_result = RequirementsResult(**req.requirements)
                questions_list, rag_contexts_by_req = analyze_build_query_for_questions(
                    build_query_obj,
                    requirements_result,
                    company_kb,
                    max_questions_per_requirement=1,
                    rag_system=rag_system,
                )
                enriched_build_query = _enrich_build_query_with_rag(
                    build_query_obj,
                    requirements_result,
                    rag_contexts_by_req,
                )
            else:
                logger.warning("Requirements not provided with build query, analyzing build query as whole")
                questions_list = analyze_build_query_for_questions_legacy(
                    build_query_obj,
                    company_kb,
                    max_questions=20,
                )
            
            all_questions = []
            questions_by_req = {}
            for idx, q in enumerate(questions_list):
                req_id = q.get("requirement_id")
                question = Question(
                    question_id=f"{req_id}-q-{idx}" if req_id else f"bq-q-{idx}",
                    requirement_id=req_id,
                    question_text=q["question_text"],
                    context=q.get("context", ""),
                    category=q.get("category", "general"),
                    priority=q.get("priority", "medium"),
                )
                all_questions.append(question.model_dump())
                
                if req_id:
                    if req_id not in questions_by_req:
                        questions_by_req[req_id] = []
                    questions_by_req[req_id].append(q)
            
            logger.info("Generated %d questions from build query", len(all_questions))
            
            response_payload: Dict[str, Any] = {
                "questions": all_questions,
                "questions_by_requirement": questions_by_req,
            }
            if req.requirements:
                response_payload["enriched_build_query"] = enriched_build_query.model_dump()
            return response_payload
        elif req.requirements:
            logger.info("Analyzing requirements for questions (legacy mode)")
            requirements_result = RequirementsResult(**req.requirements)
            
            questions_dict = analyze_requirements_for_questions(
                requirements_result.solution_requirements,
                company_kb,
                max_questions_per_requirement=1,
                rag_system=rag_system,
            )
            
            all_questions = []
            for req_id, questions in questions_dict.items():
                for q in questions:
                    question = Question(
                        question_id=f"{req_id}-q-{len(all_questions)}",
                        requirement_id=req_id,
                        question_text=q["question_text"],
                        context=q.get("context", ""),
                        category=q.get("category", "general"),
                        priority=q.get("priority", "medium"),
                    )
                    all_questions.append(question.model_dump())
            
            logger.info("Generated %d questions across %d requirements", len(all_questions), len(questions_dict))
            
            return {
                "questions": all_questions,
                "questions_by_requirement": {
                    req_id: questions
                    for req_id, questions in questions_dict.items()
                },
            }
        else:
            raise HTTPException(
                status_code=400,
                detail="Either 'build_query' or 'requirements' must be provided",
            )
    except Exception as exc:
        logger.exception("Generate questions failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail="Generate questions failed. Check server logs.",
        ) from exc

class GetNextQuestionRequest(BaseModel):
    requirements: Dict[str, Any]
    session_id: Optional[str] = None


#function to return the next critical question for iterative Q&A
@app.post("/get-next-question")
async def get_next_question_endpoint(req: GetNextQuestionRequest) -> Dict[str, Any]:
    logger.info("Get next question (session=%s)", req.session_id)
    
    try:
        company_kb = get_company_kb()
        rag_system, _ = _setup_rag_and_kb(use_rag=True)
        requirements_result = RequirementsResult(**req.requirements)
        
        previous_answers: List[Answer] = []
        rag_contexts_by_req: Dict[str, str] = {}
        if req.session_id and req.session_id in _conversation_sessions:
            context = _conversation_sessions[req.session_id]
            previous_answers = context.answers
            rag_contexts_by_req = context.rag_contexts_by_req or {}
            logger.info(
                "Session has %d previous answers and %d cached RAG contexts",
                len(previous_answers),
                len(rag_contexts_by_req),
            )
        
        question, remaining_gaps, updated_rag = get_next_critical_question(
            requirements_result=requirements_result,
            company_kb=company_kb,
            rag_system=rag_system,
            previous_answers=previous_answers,
            rag_contexts_by_req=rag_contexts_by_req,
        )
        
        if req.session_id and req.session_id in _conversation_sessions:
            _conversation_sessions[req.session_id].rag_contexts_by_req = updated_rag
        
        if question is None:
            return {
                "question": None,
                "has_more_questions": False,
                "remaining_gaps": 0,
                "message": "All critical information is available. Ready to generate response.",
            }
        
        req_id = question.get("requirement_id", "general")
        question_count = len(previous_answers)
        question["question_id"] = f"{req_id}-q-{question_count}"
        question["priority"] = "high"
        
        return {
            "question": question,
            "has_more_questions": remaining_gaps > 0,
            "remaining_gaps": remaining_gaps,
            "message": f"Please answer this question (3-5 sentences). {remaining_gaps} more question(s) may follow.",
        }
        
    except Exception as exc:
        logger.exception("Get next question failed: %s", exc)
        raise HTTPException(status_code=500, detail="Failed to get next question.") from exc


class SubmitIterativeAnswerRequest(BaseModel):
    session_id: str
    question_id: str
    question_text: str
    answer_text: str
    requirements: Dict[str, Any]


#function to submit an iterative answer and return the next question
@app.post("/submit-answer-get-next")
async def submit_answer_and_get_next(req: SubmitIterativeAnswerRequest) -> Dict[str, Any]:
    logger.info("Submit answer for %s and get next", req.question_id)
    
    if req.session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    context = _conversation_sessions[req.session_id]
    
    answer = Answer(
        question_id=req.question_id,
        question_text=req.question_text,
        answer_text=req.answer_text,
        answered_at=datetime.now().isoformat(),
    )
    context.answers.append(answer)
    logger.info("Saved answer for %s (total answers: %d)", req.question_id, len(context.answers))
    
    try:
        company_kb = get_company_kb()
        rag_system, _ = _setup_rag_and_kb(use_rag=True)
        requirements_result = RequirementsResult(**req.requirements)
        rag_contexts_by_req = context.rag_contexts_by_req or {}
        
        needs_more, next_question, remaining, updated_rag = check_if_more_questions_needed(
            requirements_result=requirements_result,
            company_kb=company_kb,
            rag_system=rag_system,
            all_answers=context.answers,
            rag_contexts_by_req=rag_contexts_by_req,
        )
        
        context.rag_contexts_by_req = updated_rag
        
        if not needs_more or next_question is None:
            return {
                "answer_saved": True,
                "next_question": None,
                "has_more_questions": False,
                "remaining_gaps": 0,
                "message": "All critical information gathered. Ready to generate response.",
            }
        
        req_id = next_question.get("requirement_id", "general")
        next_question["question_id"] = f"{req_id}-q-{len(context.answers)}"
        next_question["priority"] = "high"
        
        q_obj = Question(
            question_id=next_question["question_id"],
            requirement_id=next_question.get("requirement_id"),
            question_text=next_question["question_text"],
            context=next_question.get("context", ""),
            category=next_question.get("category", "general"),
            priority="high",
            asked_at=datetime.now().isoformat(),
        )
        context.questions.append(q_obj)
        
        return {
            "answer_saved": True,
            "next_question": next_question,
            "has_more_questions": remaining > 0,
            "remaining_gaps": remaining,
            "message": f"Answer saved. Next critical question (3-5 sentences please).",
        }
        
    except Exception as exc:
        logger.exception("Failed to get next question: %s", exc)
        return {
            "answer_saved": True,
            "next_question": None,
            "has_more_questions": False,
            "remaining_gaps": 0,
            "message": "Answer saved. Could not determine next question.",
            "error": str(exc),
        }


class CreateSessionRequest(BaseModel):
    requirement_id: Optional[str] = None


#function to create a new chat session for iterative Q&A
@app.post("/chat/session")
async def create_chat_session(req: CreateSessionRequest) -> Dict[str, Any]:
    session_id = str(uuid.uuid4())
    context = ConversationContext(
        session_id=session_id,
        requirement_id=req.requirement_id,
        created_at=datetime.now().isoformat(),
    )
    _conversation_sessions[session_id] = context
    logger.info("Created chat session: %s", session_id)
    return {"session_id": session_id}


class AddQuestionsRequest(BaseModel):
    session_id: str
    questions: List[Dict[str, Any]]


#function to add questions to an existing chat session
@app.post("/chat/questions")
async def add_questions(req: AddQuestionsRequest) -> Dict[str, Any]:
    if req.session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    context = _conversation_sessions[req.session_id]
    
    for q_dict in req.questions:
        question = Question(**q_dict)
        question.asked_at = datetime.now().isoformat()
        context.questions.append(question)
    
    logger.info("Added %d questions to session %s", len(req.questions), req.session_id)
    return {"status": "ok", "questions_count": len(context.questions)}


class SubmitAnswerRequest(BaseModel):
    session_id: str
    question_id: str
    answer_text: str


#function to submit an answer to a question in a chat session
@app.post("/chat/answer")
async def submit_answer(req: SubmitAnswerRequest) -> Dict[str, Any]:
    if req.session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    context = _conversation_sessions[req.session_id]
    
    question = next((q for q in context.questions if q.question_id == req.question_id), None)
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    answer = Answer(
        question_id=req.question_id,
        answer_text=req.answer_text,
        answered_at=datetime.now().isoformat(),
    )
    context.answers.append(answer)
    question.answered = True
    
    remaining_questions = [q for q in context.questions if not q.answered]
    auto_resolved_ids: list[str] = []
    if remaining_questions:
        try:
            auto_resolved_ids = infer_answered_questions_from_answer(
                answered_question=question,
                answer_text=req.answer_text,
                remaining_questions=remaining_questions,
            )
        except Exception as exc:
            logger.exception(
                "Failed to infer additionally answered questions for session %s: %s",
                req.session_id,
                exc,
            )
            auto_resolved_ids = []
    
    for qid in auto_resolved_ids:
        extra_q = next((q for q in context.questions if q.question_id == qid), None)
        if extra_q and not extra_q.answered:
            extra_answer = Answer(
                question_id=qid,
                answer_text=req.answer_text,
                answered_at=datetime.now().isoformat(),
            )
            context.answers.append(extra_answer)
            extra_q.answered = True
    
    if auto_resolved_ids:
        logger.info(
            "Answer to question %s in session %s also resolved %d other question(s): %s",
            req.question_id,
            req.session_id,
            len(auto_resolved_ids),
            auto_resolved_ids,
        )
    else:
        logger.info(
            "Answer submitted for question %s in session %s (no additional questions auto-resolved)",
            req.question_id,
            req.session_id,
        )
    
    return {
        "status": "ok",
        "answer": answer.model_dump(),
        "auto_resolved_question_ids": auto_resolved_ids,
    }


#function to return the state of a chat session (questions and answers)
@app.get("/chat/session/{session_id}")
async def get_session(session_id: str) -> Dict[str, Any]:
    if session_id not in _conversation_sessions:
        raise HTTPException(status_code=404, detail="Session not found")
    
    context = _conversation_sessions[session_id]
    return {
        "session_id": context.session_id,
        "requirement_id": context.requirement_id,
        "questions": [q.model_dump() for q in context.questions],
        "answers": [a.model_dump() for a in context.answers],
        "created_at": context.created_at,
        "qa_context": context.get_qa_context(),
    }


#function to enrich build query with session context
@app.post("/enrich-build-query")
async def enrich_build_query_endpoint(req: EnrichBuildQueryRequest) -> Dict[str, Any]:
    try:
        build_query = BuildQuery(**req.build_query)
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail=f"Invalid build_query payload: {str(exc)}",
        ) from exc

    if req.session_id and req.session_id in _conversation_sessions:
        context = _conversation_sessions[req.session_id]
        qa_context = context.get_qa_context()
        if qa_context:
            original_text = build_query.query_text or ""
            if original_text and qa_context:
                enriched_text = f"{original_text}\n\n--- Q&A Context from Session ---\n{qa_context}"
                build_query.query_text = enriched_text
                logger.info(
                    "Enriched build query with Q&A context from session %s (%d answers, added %d chars)",
                    req.session_id,
                    len(context.answers),
                    len(qa_context),
                )
            else:
                logger.info(
                    "Enrich build query called for session %s (answers=%d) – no Q&A context to add",
                    req.session_id,
                    len(context.answers),
                )
        else:
            logger.info(
                "Enrich build query called for session %s (answers=%d) – Q&A context is empty",
                req.session_id,
                len(context.answers),
            )
    else:
        logger.info("Enrich build query called without valid session_id; returning original build query text")

    return build_query.model_dump()


_response_cache: Dict[str, List[Dict[str, Any]]] = {}


class PreviewResponseRequest(BaseModel):
    preprocess: Dict[str, Any]
    requirements: Dict[str, Any]
    use_rag: bool = True
    num_retrieval_chunks: int = 5
    session_id: Optional[str] = None


@app.post("/preview-responses")
async def preview_responses_endpoint(req: PreviewResponseRequest) -> Dict[str, Any]:
    logger.info("Preview responses endpoint called")
    try:
        from backend.models import RequirementsResult
        
        preprocess_result = PreprocessResult(**req.preprocess)
        extraction_result = _extraction_from_preprocess(preprocess_result)
        requirements_result = RequirementsResult(**req.requirements)
        
        if not requirements_result.solution_requirements:
            raise HTTPException(
                status_code=400,
                detail="No solution requirements found. Cannot generate responses.",
            )
        
        if requirements_result.structure_detection is None:
            structure_detection_dict = detect_structure(requirements_result.response_structure_requirements)
            requirements_result.structure_detection = StructureDetectionResult(**structure_detection_dict)
        
        structure_detection = requirements_result.structure_detection
        
        rag_system, knowledge_base = _setup_rag_and_kb(req.use_rag)
        
        qa_context = ""
        if req.session_id and req.session_id in _conversation_sessions:
            context = _conversation_sessions[req.session_id]
            qa_context = context.get_qa_context()
        
        validation_errors = validate_before_generation(extraction_result, requirements_result)
        if validation_errors:
            raise HTTPException(status_code=400, detail="Validation failed: " + "; ".join(validation_errors))
        
        individual_responses = []
        total_requirements = len(requirements_result.solution_requirements)
        
        for idx, solution_req in enumerate(requirements_result.solution_requirements, 1):
            key_phrase = _extract_key_phrase(solution_req.source_text)
            
            try:
                build_query_obj = build_query_for_single_requirement(
                    extraction_result=extraction_result,
                    single_requirement=solution_req,
                    all_response_structure_requirements=requirements_result.response_structure_requirements,
                )
                build_query_obj.confirmed = True
                
                result = run_response_agent(
                    build_query=build_query_obj,
                    knowledge_base=knowledge_base,
                    qa_context=qa_context,
                )
                
                quality_assessment = assess_response_quality(solution_req, result.response_text)
                
                individual_responses.append({
                    "requirement_id": solution_req.id,
                    "requirement_text": solution_req.source_text,
                    "key_phrase": key_phrase,
                    "response": result.response_text,
                    "notes": result.notes,
                    "quality": quality_assessment,
                })
            except Exception as req_exc:
                logger.error("Failed to generate response for requirement %s: %s", solution_req.id, req_exc)
                individual_responses.append({
                    "requirement_id": solution_req.id,
                    "requirement_text": solution_req.source_text,
                    "key_phrase": key_phrase,
                    "response": f"[ERROR: Failed to generate response: {str(req_exc)}]",
                    "notes": f"Error: {str(req_exc)}",
                    "quality": {
                        "score": 0.0,
                        "completeness": "incomplete",
                        "relevance": "low",
                        "issues": [f"Generation failed: {str(req_exc)}"],
                        "suggestions": ["Fix the error and regenerate"],
                    },
                })
        
        preview_id = str(uuid.uuid4())
        _response_cache[preview_id] = individual_responses
        
        return {
            "preview_id": preview_id,
            "responses": individual_responses,
            "total": len(individual_responses),
        }
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Preview responses failed: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Preview generation failed: {str(exc)}",
        ) from exc


class PreviewContextRequest(BaseModel):
    preprocess: Dict[str, Any]
    requirements: Dict[str, Any]
    use_rag: bool = True
    num_retrieval_chunks: int = 5
    session_id: Optional[str] = None


@app.post("/preview-context")
@app.get("/preview-context")
async def preview_context_endpoint(req: PreviewContextRequest | None = None) -> Dict[str, Any]:
    if req is None:
        return {
            "message": "POST /preview-context expected. Use POST with JSON body {preprocess, requirements, use_rag, num_retrieval_chunks, session_id}.",
            "note": "This GET response is a diagnostic fallback. Use POST to fetch RAG context snippets.",
        }
    logger.info("Preview context endpoint called")
    try:
        preprocess_result = PreprocessResult(**req.preprocess)
        extraction_result = _extraction_from_preprocess(preprocess_result)
        requirements_result = RequirementsResult(**req.requirements)

        rag_system, knowledge_base = _setup_rag_and_kb(req.use_rag)

        rag_contexts_by_req: Dict[str, List[Dict[str, Any]]] = {}
        if rag_system:
            for solution_req in requirements_result.solution_requirements:
                qtext = solution_req.source_text or ""
                try:
                    results = rag_system.search(qtext, k=req.num_retrieval_chunks)
                except Exception as e:
                    logger.exception("RAG search failed for requirement %s: %s", solution_req.id, e)
                    results = []
                rag_contexts_by_req[solution_req.id] = results
        else:
            for solution_req in requirements_result.solution_requirements:
                rag_contexts_by_req[solution_req.id] = []

        session_data = None
        if req.session_id and req.session_id in _conversation_sessions:
            session = _conversation_sessions[req.session_id]
            session_data = session.model_dump()

        return {
            "rag_contexts_by_requirement": rag_contexts_by_req,
            "session": session_data,
        }
    except Exception as exc:
        logger.exception("Preview context failed: %s", exc)
        raise HTTPException(status_code=500, detail=f"Preview context failed: {str(exc)}") from exc


class UpdateResponseRequest(BaseModel):
    preview_id: str
    requirement_id: str
    response_text: str


@app.post("/update-response")
async def update_response_endpoint(req: UpdateResponseRequest) -> Dict[str, Any]:
    if req.preview_id not in _response_cache:
        raise HTTPException(status_code=404, detail="Preview not found")
    
    responses = _response_cache[req.preview_id]
    response = next((r for r in responses if r["requirement_id"] == req.requirement_id), None)
    
    if not response:
        raise HTTPException(status_code=404, detail="Requirement not found in preview")
    
    response["response"] = req.response_text
    logger.info("Updated response for requirement %s in preview %s", req.requirement_id, req.preview_id)
    
    return {"status": "ok", "updated": True}


class GeneratePDFFromPreviewRequest(BaseModel):
    preview_id: str
    preprocess: Dict[str, Any]
    requirements: Dict[str, Any]
    format: str = "pdf" 


@app.post("/generate-pdf-from-preview")
async def generate_pdf_from_preview_endpoint(req: GeneratePDFFromPreviewRequest) -> Response:
    if req.preview_id not in _response_cache:
        raise HTTPException(status_code=404, detail="Preview not found")
    
    individual_responses = _response_cache[req.preview_id]
    
    preprocess_result = PreprocessResult(**req.preprocess)
    extraction_result = _extraction_from_preprocess(preprocess_result)
    requirements_result = RequirementsResult(**req.requirements)
    
    rfp_title = _extract_title_from_key_requirements(extraction_result.key_requirements_summary)
    
    project_root = Path(__file__).parent.parent
    timestamp = int(time.time())
    
    logger.info(
        "Generate PDF from preview called (format=%s, preview_id=%s)",
        req.format,
        req.preview_id,
    )

    if req.format == "docx":
        try:
            from backend.document_formatter import generate_rfp_docx
        except ImportError as import_err:
            logger.error("DOCX generation not available: %s", import_err)
            raise HTTPException(status_code=500, detail="DOCX generation not available. Install python-docx.") from import_err
        
        output_dir = project_root / "output" / "docx"
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"rfp_response_{extraction_result.language}_{timestamp}.docx"
        output_path = output_dir / filename
        logger.info(
            "Generating DOCX from preview %s into %s",
            req.preview_id,
            output_path,
        )

        docx_bytes = generate_rfp_docx(
            individual_responses=individual_responses,
            requirements_result=requirements_result,
            extraction_result=extraction_result,
            rfp_title=rfp_title,
            output_path=output_path,
        )
        
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(docx_bytes)),
            }
        )
    
    elif req.format == "markdown":
        from backend.document_formatter import generate_rfp_markdown
        
        output_dir = project_root / "output" / "markdown"
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"rfp_response_{extraction_result.language}_{timestamp}.md"
        output_path = output_dir / filename
        logger.info(
            "Generating Markdown from preview %s into %s",
            req.preview_id,
            output_path,
        )

        markdown_bytes = generate_rfp_markdown(
            individual_responses=individual_responses,
            requirements_result=requirements_result,
            extraction_result=extraction_result,
            rfp_title=rfp_title,
            output_path=output_path,
        )
        
        return Response(
            content=markdown_bytes,
            media_type="text/markdown",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(markdown_bytes)),
            }
        )
    
    else:
        from backend.document_formatter import generate_rfp_pdf
        
        output_dir = project_root / "output" / "pdfs"
        output_dir.mkdir(parents=True, exist_ok=True)
        filename = f"rfp_response_{extraction_result.language}_{timestamp}.pdf"
        pdf_path = output_dir / filename
        logger.info(
            "Generating PDF from preview %s (in memory)",
            req.preview_id,
        )

        pdf_bytes = generate_rfp_pdf(
            individual_responses=individual_responses,
            requirements_result=requirements_result,
            extraction_result=extraction_result,
            rfp_title=rfp_title,
            output_path=None,
        )
        
        try:
            pdf_path.write_bytes(pdf_bytes)
            logger.info("PDF also saved to disk: %s", pdf_path.absolute())
        except Exception as save_err:
            logger.warning("Failed to save PDF to disk: %s (continuing with in-memory response)", save_err)
        
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(len(pdf_bytes)),
                "X-PDF-Path": str(pdf_path.absolute()),
            }
        )


#function to convert HTML content to DOCX using python-docx
def _html_to_docx(html_content: str) -> bytes:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        raise ImportError("beautifulsoup4 is not installed. Install it with: pip install beautifulsoup4")
    
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
    except ImportError:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    
    doc = Document()
    
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'ul', 'ol', 'table', 'img', 'br', 'div']):
        if element.name == 'p':
            para = doc.add_paragraph()
            _add_text_to_paragraph(para, element)
        elif element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            heading = doc.add_heading(level=min(level, 9))
            _add_text_to_paragraph(heading, element)
        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                para = doc.add_paragraph(style='List Bullet' if element.name == 'ul' else 'List Number')
                _add_text_to_paragraph(para, li)
        elif element.name == 'table':
            _add_table_to_doc(doc, element)
        elif element.name == 'img':
            src = element.get('src', '')
            if src.startswith('data:image'):
                try:
                    import base64
                    header, encoded = src.split(',', 1)
                    img_data = base64.b64decode(encoded)
                    from io import BytesIO
                    doc.add_picture(BytesIO(img_data))
                except Exception as e:
                    logger.warning(f"Failed to add image: {e}")
        elif element.name == 'br':
            doc.add_paragraph()
        elif element.name == 'div' and element.get_text(strip=True):
            para = doc.add_paragraph()
            _add_text_to_paragraph(para, element)
    
    if len(doc.paragraphs) == 0:
        text_content = soup.get_text()
        for line in text_content.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
    
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


#function to add text content to a paragraph, handling formatting
def _add_text_to_paragraph(para, element):
    from bs4 import NavigableString
    
    for content in element.contents:
        if isinstance(content, NavigableString):
            text = str(content).strip()
            if text:
                para.add_run(text)
        elif hasattr(content, 'name'):
            if content.name == 'strong' or content.name == 'b':
                run = para.add_run(content.get_text())
                run.bold = True
            elif content.name == 'em' or content.name == 'i':
                run = para.add_run(content.get_text())
                run.italic = True
            elif content.name == 'u':
                run = para.add_run(content.get_text())
                run.underline = True
            else:
                para.add_run(content.get_text())


#function to add a table from HTML to DOCX
def _add_table_to_doc(doc, table_element):
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    max_cols = max(len(row.find_all(['td', 'th'])) for row in rows)
    if max_cols == 0:
        return
    
    table = doc.add_table(rows=len(rows), cols=max_cols)
    
    for row_idx, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for col_idx, cell in enumerate(cells[:max_cols]):
            table.rows[row_idx].cells[col_idx].text = cell.get_text(strip=True)
            # Style header cells
            if cell.name == 'th':
                for para in table.rows[row_idx].cells[col_idx].paragraphs:
                    for run in para.runs:
                        run.bold = True


class SaveDocxRequest(BaseModel):
    docx_bytes: Optional[str] = None
    html_content: Optional[str] = None
    filename: Optional[str] = None

@app.post("/save-docx")
#function to save a DOCX file to the output folder (from blob or HTML)
async def save_docx_endpoint(req: SaveDocxRequest) -> Dict[str, Any]:
    import base64
    
    try:

        if req.html_content:
            logger.info("Converting HTML content to DOCX...")
            docx_bytes = _html_to_docx(req.html_content)
            logger.info("HTML converted to DOCX: %d bytes", len(docx_bytes))
        elif req.docx_bytes:
            docx_bytes = base64.b64decode(req.docx_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail="Either 'docx_bytes' or 'html_content' must be provided",
            )
        
        project_root = Path(__file__).parent.parent
        output_dir = project_root / "output" / "docx"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        if not req.filename:
            timestamp = int(time.time())
            req.filename = f"rfp_response_en_{timestamp}.docx"
        
        if not req.filename.endswith('.docx'):
            req.filename += '.docx'
        
        output_path = output_dir / req.filename
        
        output_path.write_bytes(docx_bytes)
        
        logger.info("DOCX saved successfully: %d bytes to %s", len(docx_bytes), output_path.absolute())
        
        return {
            "status": "ok",
            "filename": req.filename,
            "path": str(output_path.absolute()),
            "size": len(docx_bytes),
        }
    except Exception as exc:
        logger.exception("Failed to save DOCX: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Failed to save DOCX: {str(exc)}",
        ) from exc


@app.get("/health")
async def health() -> Dict[str, Any]:
    return {"status": "ok"}


@app.get("/{path:path}")
async def serve_frontend(path: str):
    if path.startswith(("assets/", "src/", "public/")):
        raise HTTPException(status_code=404, detail="Static file not found")
    
    project_root = Path(__file__).resolve().parent.parent
    index_path = project_root / "frontend" / "dist" / "index.html"
    if not index_path.exists():
        index_path = project_root / "frontend" / "index.html"
    if not index_path.exists():
        raise HTTPException(status_code=404, detail="Frontend not found")
    return HTMLResponse(index_path.read_text(encoding="utf-8"))


