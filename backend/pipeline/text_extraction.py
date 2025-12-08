from __future__ import annotations

import base64
import io
import logging
from pathlib import Path
from typing import List, Union
import tempfile


from PIL import Image  # type: ignore

from backend.llm.client import chat_completion_with_vision


logger = logging.getLogger(__name__)

PathLike = Union[str, Path]

TEXT_EXTRACTION_MODEL = "Qwen/Qwen2.5-VL-7B-Instruct"

MIN_TEXT_LENGTH = 100


def _pdf_to_images(pdf_path: Path) -> List[Image.Image]:
    try:
        from pdf2image import convert_from_path  # type: ignore
    except ImportError:
        raise ImportError(
            "pdf2image is required for PDF text extraction. "
            "Install with: pip install pdf2image. "
            "You may also need poppler-utils: brew install poppler (macOS) or apt-get install poppler-utils (Linux)"
        )

    logger.info("[Vision OCR Fallback] Converting PDF to images: %s", pdf_path.name)
    logger.info("[Vision OCR Fallback] Using DPI: 200 (for OCR quality)")
    images = convert_from_path(str(pdf_path), dpi=200)
    logger.info("[Vision OCR Fallback] Converted PDF to %d page image(s)", len(images))
    return images

def _docx_to_images(docx_path: Path) -> List[Image.Image]:
    try:
        from docx2pdf import convert  # type: ignore
        from pdf2image import convert_from_path  # type: ignore
    except ImportError:
        raise ImportError(
            "docx2pdf and pdf2image are required for DOCX text extraction. "
            "Install with: pip install docx2pdf pdf2image. "
            "You may also need poppler-utils: brew install poppler (macOS) or apt-get install poppler-utils (Linux)"
        )

    logger.info("[Vision OCR Fallback] Converting DOCX to images: %s", docx_path.name)
    logger.info("[Vision OCR Fallback] Step 1: Converting DOCX to PDF...")
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf:
        tmp_pdf_path = Path(tmp_pdf.name)
        try:
            convert(str(docx_path), str(tmp_pdf_path))
            logger.info("[Vision OCR Fallback] Step 2: Converting PDF to images (DPI: 200)...")
            images = convert_from_path(str(tmp_pdf_path), dpi=200)
            logger.info("[Vision OCR Fallback] Converted DOCX to %d page image(s)", len(images))
            return images
        finally:
            try:
                tmp_pdf_path.unlink(missing_ok=True)
                logger.debug("[Vision OCR Fallback] Cleaned up temporary PDF file")
            except Exception:
                pass

def _image_to_base64(image: Image.Image) -> str:
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
    return f"data:image/png;base64,{img_base64}"

def _extract_text_from_images(images: List[Image.Image]) -> str:
    logger.info("[Vision OCR] Starting OCR extraction from %d image(s) using model: %s", 
               len(images), TEXT_EXTRACTION_MODEL)
    logger.info("[Vision OCR] This may take a while depending on document size...")
    all_text_parts: List[str] = []
    for idx, image in enumerate(images, 1):
        logger.info("[Vision OCR] Processing page %d/%d...", idx, len(images))
        img_base64 = _image_to_base64(image)
        logger.debug("[Vision OCR] Page %d: encoded image to base64 (%d bytes)", 
                    idx, len(img_base64))

        prompt = (
            "Extract ALL text from this document page. "
            "Preserve the structure, formatting, and order of the text as it appears. "
            "Include all headings, paragraphs, lists, tables (as text), and any other textual content. "
            "Do not summarize or skip any text. Return the complete text content of this page."
        )

        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": img_base64},
                    },
                ],
            }
        ]

        logger.info("[Vision OCR] Sending page %d/%d to vision model API...", idx, len(images))
        page_text = chat_completion_with_vision(
            model=TEXT_EXTRACTION_MODEL,
            messages=messages,
            temperature=0.0,
            max_tokens=None,
        )

        all_text_parts.append(page_text)
        logger.info("[Vision OCR] Page %d/%d: extracted %d characters", idx, len(images), len(page_text))
    full_text = "\n\n--- Page Break ---\n\n".join(all_text_parts)
    logger.info("[Vision OCR] OCR extraction completed: %d total characters from %d page(s)", 
               len(full_text), len(images))
    return full_text

def _extract_text_from_pdf_direct(pdf_path: Path) -> str:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        logger.warning("[PDF Direct Extraction] pdfplumber library not available, skipping direct extraction")
        return ""
    try:
        logger.info("[PDF Direct Extraction] Starting direct text extraction from PDF: %s", pdf_path.name)
        text_parts: List[str] = []
        total_pages = 0       
        with pdfplumber.open(str(pdf_path)) as pdf:
            total_pages = len(pdf.pages)
            logger.info("[PDF Direct Extraction] PDF has %d page(s)", total_pages)          
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
                    logger.info("[PDF Direct Extraction] Page %d/%d: extracted %d characters", 
                               page_num, total_pages, len(page_text))
                else:
                    logger.warning("[PDF Direct Extraction] Page %d/%d: no text found (may be image-only)", 
                                 page_num, total_pages)   
        full_text = "\n\n--- Page Break ---\n\n".join(text_parts)
        pages_with_text = len(text_parts)
        logger.info("[PDF Direct Extraction] Completed: %d characters from %d/%d pages with text", 
                   len(full_text), pages_with_text, total_pages)
        return full_text
    except Exception as e:
        logger.error("[PDF Direct Extraction] Failed with error: %s", str(e), exc_info=True)
        return ""


def _extract_text_from_docx_direct(docx_path: Path) -> str:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        logger.warning("[DOCX Direct Extraction] python-docx library not available, skipping direct extraction")
        return ""
    try:
        logger.info("[DOCX Direct Extraction] Starting direct text extraction from DOCX: %s", docx_path.name)
        doc = Document(str(docx_path))     
        text_parts: List[str] = []
        para_count = 0
        table_count = 0     
        logger.info("[DOCX Direct Extraction] Extracting paragraphs...")
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                para_count += 1     
        logger.info("[DOCX Direct Extraction] Found %d paragraphs with text", para_count)
        logger.info("[DOCX Direct Extraction] Extracting tables...")
        for table_idx, table in enumerate(doc.tables, 1):
            table_text_parts: List[str] = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    table_text_parts.append(row_text)
            if table_text_parts:
                text_parts.append("\n".join(table_text_parts))
                table_count += 1
                logger.info("[DOCX Direct Extraction] Table %d: extracted %d rows", table_idx, len(table_text_parts))
        full_text = "\n\n".join(text_parts)
        logger.info("[DOCX Direct Extraction] Completed: %d characters from %d paragraphs and %d tables", 
                   len(full_text), para_count, table_count)
        return full_text
    except Exception as e:
        logger.error("[DOCX Direct Extraction] Failed with error: %s", str(e), exc_info=True)
        return ""


def extract_text_from_file(path: PathLike) -> str:
    p = Path(path)
    suffix = p.suffix.lower()
    logger.info("=" * 80)
    logger.info("[Text Extraction] Starting text extraction from file: %s", p.name)
    logger.info("[Text Extraction] File type: %s", suffix.upper())
    logger.info("=" * 80)
    direct_text = ""
    extraction_method = ""
    if suffix == ".pdf":
        logger.info("[Text Extraction] Attempting direct PDF extraction (pdfplumber)...")
        direct_text = _extract_text_from_pdf_direct(p)
        extraction_method = "pdfplumber"
    elif suffix in {".docx", ".doc"}:
        if suffix == ".docx":
            logger.info("[Text Extraction] Attempting direct DOCX extraction (python-docx)...")
            direct_text = _extract_text_from_docx_direct(p)
            extraction_method = "python-docx"
        else:
            logger.info("[Text Extraction] .doc files not supported for direct extraction, skipping to OCR")
    else:
        raise ValueError(f"Unsupported file type for path: {path}")
    direct_text_length = len(direct_text.strip()) if direct_text else 0
    if direct_text and direct_text_length >= MIN_TEXT_LENGTH:
        logger.info("=" * 80)
        logger.info("[Text Extraction] ✓ SUCCESS: Direct extraction successful!")
        logger.info("[Text Extraction] Method: %s", extraction_method)
        logger.info("[Text Extraction] Extracted: %d characters", direct_text_length)
        logger.info("[Text Extraction] Using direct extraction result (no OCR needed)")
        logger.info("=" * 80)
        return direct_text
    logger.info("=" * 80)
    if direct_text:
        logger.warning(
            "[Text Extraction] Direct extraction returned only %d characters (minimum required: %d)",
            direct_text_length,
            MIN_TEXT_LENGTH
        )
        logger.warning("[Text Extraction] Document may be image-based or scanned. Falling back to vision model OCR...")
    else:
        logger.warning("[Text Extraction] Direct extraction failed or returned no text.")
        logger.warning("[Text Extraction] Falling back to vision model OCR (this will take longer)...")
    logger.info("=" * 80)
    logger.info("[Text Extraction] Preparing document for OCR...")
    if suffix == ".pdf":
        images = _pdf_to_images(p)
    elif suffix in {".docx", ".doc"}:
        images = _docx_to_images(p)
    else:
        raise ValueError(f"Unsupported file type for path: {path}")
    if not images:
        logger.error("[Text Extraction] ✗ ERROR: No images generated from file: %s", path)
        logger.warning("[Text Extraction] Returning any text from direct extraction (if available)")
        # If we have some text from direct extraction, return it even if short
        result = direct_text if direct_text else ""
        logger.info("[Text Extraction] Final result: %d characters", len(result))
        return result
    ocr_text = _extract_text_from_images(images)
    logger.info("=" * 80)
    logger.info("[Text Extraction] ✓ COMPLETED: OCR extraction finished")
    logger.info("[Text Extraction] Method: Vision model OCR (%s)", TEXT_EXTRACTION_MODEL)
    logger.info("[Text Extraction] Extracted: %d characters from %d page(s)", len(ocr_text), len(images))
    logger.info("=" * 80)
    return ocr_text


